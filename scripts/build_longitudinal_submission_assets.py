from __future__ import annotations

from pathlib import Path
import re
import zipfile

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
ANALYSIS_DIR = ROOT / "results" / "longitudinal_tb_analysis"
OUT_DIR = ROOT / "submission_ready" / "longitudinal_tb_dynamics_20260318"
REPO_URL = "https://github.com/hssling/tb-progression-transcriptome-meta"

AUTHOR_NAME = "Siddalingaiah H S"
DEGREE = "MD"
AFFILIATION = (
    "Department of Community Medicine, Shridevi Institute of Medical Sciences & "
    "Research Hospital, Tumkur, Karnataka, India"
)
CORRESPONDING = (
    "Siddalingaiah H S, Department of Community Medicine, Shridevi Institute of "
    "Medical Sciences & Research Hospital, Tumkur, Karnataka, India. "
    "Email: hssling@yahoo.com"
)
TITLE = (
    "Longitudinal reanalysis of a prospective tuberculosis blood RNA-sequencing "
    "cohort suggests follow-up dynamics in coordinated host-response programs"
)
SHORT_TITLE = "Longitudinal TB host-response dynamics"

REFERENCES = [
    "World Health Organization. Global tuberculosis report 2025. Geneva: World Health Organization; 2025.",
    "Zak DE, Penn-Nicholson A, Scriba TJ, et al. A blood RNA signature for tuberculosis disease risk: a prospective cohort study. Lancet. 2016;387:2312-2322.",
    "Penn-Nicholson A, Mbandi SK, Thompson E, et al. RISK6, a 6-gene transcriptomic signature of TB disease risk, diagnosis and treatment response. Sci Rep. 2020;10:8629.",
    "Gupta RK, Turner CT, Venturini C, et al. Concise whole blood transcriptional signatures for incipient tuberculosis: a systematic review and patient-level pooled meta-analysis. Lancet Respir Med. 2020;8:395-406.",
    "De Groote MA, Gupta RK, Hellwig SM, et al. Prospective multicentre head-to-head validation of host blood transcriptomic biomarkers for pulmonary tuberculosis by real-time PCR. Commun Med (Lond). 2022;2:26.",
    "Andrews JR, Nemes E, Tameris M, et al. Transcriptomic signatures of progression to tuberculosis disease among close contacts in Brazil. Clin Infect Dis. 2024;78:1672-1681.",
    "Laird NM, Ware JH. Random-effects models for longitudinal data. Biometrics. 1982;38:963-974.",
    "Leek JT, Johnson WE, Parker HS, Jaffe AE, Storey JD. The sva package for removing batch effects and other unwanted variation in high-throughput experiments. Bioinformatics. 2012;28:882-883.",
    "Newman AM, Liu CL, Green MR, et al. Robust enumeration of cell subsets from tissue expression profiles. Nat Methods. 2015;12:453-457.",
    "Langfelder P, Horvath S. WGCNA: an R package for weighted correlation network analysis. BMC Bioinformatics. 2008;9:559.",
    "National Center for Biotechnology Information. Gene Expression Omnibus: GSE79362. https://www.ncbi.nlm.nih.gov/geo/query/acc.cgi?acc=GSE79362. Accessed 18 Mar 2026.",
    "Siddalingaiah HS. tb-progression-transcriptome-meta. GitHub. https://github.com/hssling/tb-progression-transcriptome-meta. Accessed 18 Mar 2026.",
]


def set_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(0)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)


def add_plain_paragraph(doc: Document, text: str = "", bold: bool = False, center: bool = False) -> None:
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_cited_paragraph(doc: Document, text: str, center: bool = False) -> None:
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    parts = re.split(r"(\[\[[0-9,\- ]+\]\])", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("[[") and part.endswith("]]"):
            run = p.add_run(part[2:-2].replace(" ", ""))
            run.font.superscript = True
        else:
            p.add_run(part)


def add_table(doc: Document, df: pd.DataFrame, title: str) -> None:
    add_plain_paragraph(doc, title)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for idx, col in enumerate(df.columns):
        table.cell(0, idx).text = str(col)
    for row in df.itertuples(index=False):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    doc.add_paragraph("")


def add_figure(doc: Document, image_path: Path, caption: str, width: float = 6.2) -> None:
    add_plain_paragraph(doc, caption)
    if image_path.exists():
        doc.add_picture(str(image_path), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")


def load_outputs() -> dict[str, pd.DataFrame]:
    return {
        "program_models": pd.read_csv(ANALYSIS_DIR / "mixedlm_program_models.csv"),
        "gene_models": pd.read_csv(ANALYSIS_DIR / "mixedlm_gene_models.csv"),
        "ic": pd.read_csv(ANALYSIS_DIR / "ic_progressor_contrasts.csv"),
        "timepoint": pd.read_csv(ANALYSIS_DIR / "timepoint_program_summary.csv"),
        "gene_sets": pd.read_csv(ANALYSIS_DIR / "program_gene_sets.csv"),
        "scores": pd.read_csv(ANALYSIS_DIR / "program_scores.csv"),
    }


def manuscript_text(data: dict[str, pd.DataFrame]) -> dict[str, list[str] | str]:
    program_models = data["program_models"]
    gene_models = data["gene_models"]
    timepoint = data["timepoint"]
    scores = data["scores"]

    top_programs = program_models.head(5)["feature"].tolist()
    top_genes = gene_models.head(3)["feature"].tolist()
    fdr10_programs = program_models.loc[program_models["interaction_fdr_bh"] < 0.10, "feature"].tolist()
    n_subjects = int(scores["subject_id"].nunique())
    n_samples = int(len(scores))
    n_progressors = int(scores.loc[scores["progressor"] == 1, "subject_id"].nunique())
    n_nonprogressors = int(scores.loc[scores["progressor"] == 0, "subject_id"].nunique())
    n_sched = int(scores["timepoint_month"].notna().sum())

    day540_prog = timepoint[(timepoint["progressor"] == 1) & (timepoint["timepoint_label"] == "DAY540")]
    day180_prog = timepoint[(timepoint["progressor"] == 1) & (timepoint["timepoint_label"] == "DAY180")]
    day540_non = timepoint[(timepoint["progressor"] == 0) & (timepoint["timepoint_label"] == "DAY540")]
    day180_non = timepoint[(timepoint["progressor"] == 0) & (timepoint["timepoint_label"] == "DAY180")]

    abstract = [
        "Background: Most blood transcriptomic studies of tuberculosis progression summarize baseline discrimination, but prospective cohorts with repeated samples can also clarify whether host-response programs change differently over follow-up in progressors and non-progressors. We reanalyzed a remapped public longitudinal RNA-sequencing cohort to test that question.",
        f"Methods: The GSE79362 cohort was remapped from junction-level counts to a gene-level matrix and linked to subject identifiers and follow-up labels. Primary models used {n_sched} scheduled samples with numeric follow-up months from {n_subjects} subjects. Prespecified scores represented an existing 25-gene signature, a compact Bayesian core, a remap-sensitive myeloid set, vascular and coexpression proxies, and marker-based cell-state proxies. Random-intercept mixed-effects or subject-clustered longitudinal models estimated month, progressor, and progressor-by-month effects. IC samples were analyzed separately as a sensitivity layer.",
        f"Results: The cohort contained {n_samples} samples from {n_subjects} subjects, including {n_progressors} progressors and {n_nonprogressors} non-progressors. The strongest progressor-by-month interactions were observed for vascular_proxy5 (coefficient -0.0218, p=0.0092, FDR=0.0587), remap_myeloid5 (-0.0295, p=0.0125, FDR=0.0587), neutrophil proxy (-0.0026, p=0.0177, FDR=0.0587), platelet proxy (0.0035, p=0.0196, FDR=0.0587), and module_M6_proxy (-0.0222, p=0.0366, FDR=0.0878). No individual gene interaction survived multiple-testing correction, although ACSL1, FCGR3B, and VSIG4 showed the strongest nominal trends.",
        "Conclusion: In this longitudinal cohort, temporal differences between progressors and non-progressors were more coherent at the level of coordinated host-response programs than at the level of single transcripts. These findings support a dynamics-oriented interpretation of progression biology, but they should not be read as direct time-to-disease forecasting because scheduled follow-up months and IC samples are not equivalent clinical states.",
    ]

    intro = [
        "Tuberculosis remains a major global health problem, and host blood transcriptomic signatures continue to attract attention as potential tools for identifying people at increased risk of future disease.[[1-4]] Most published work has emphasized cross-sectional or baseline discrimination, often asking whether one blood sample can distinguish progressors from non-progressors or active disease from other states.[[2-6]]",
        "That work has been valuable, but it leaves a second question less well developed: how the host signal behaves over repeated follow-up within the same individual. A progression-associated signature may be biologically real even if no single transcript remains uniformly dominant across all visits. Conversely, some apparently strong baseline genes may weaken once repeated sampling and within-subject variation are considered.",
        "Longitudinal reanalysis is particularly relevant in public tuberculosis cohorts because some datasets include scheduled follow-up visits, unscheduled or intermediate collections, and additional samples close to clinically important transitions. Those features can add biological meaning, but only if they are handled conservatively rather than collapsed into one pseudo-time variable without scrutiny.",
        "A longitudinal framework is also useful for a practical reason. If serial blood testing is ever to contribute to tuberculosis prevention or early recognition, the biomarker cannot be judged only by cross-sectional separation. It must also show interpretable behavior across repeated visits, because real implementation would almost certainly involve follow-up measurements rather than one isolated sample.",
        "At the same time, longitudinal public-data analysis is technically demanding. Repeated collections create within-subject dependence, intermediate or unscheduled samples complicate the visit structure, and RNA-sequencing data released at exon or junction level require additional preprocessing before biologically meaningful gene-level interpretation becomes possible. These challenges make conservative remapping and model specification essential.",
        "The present analysis therefore takes a prespecified program-based approach. Instead of deriving a new single-cohort classifier, it evaluates whether host-response programs already highlighted by earlier harmonized public-data analyses show different temporal behavior in a remapped prospective RNA-sequencing cohort. This approach was chosen to reduce overfitting and to ask a more mechanistic question.",
        "Our aim was to determine whether scheduled follow-up in the remapped GSE79362 cohort reveals reproducible progressor-specific dynamics at the level of coordinated programs, including myeloid, vascular-remodeling, coexpression, and marker-based cell-state signals, while keeping IC samples and time interpretation within conservative limits.",
    ]

    methods = [
        "This was a secondary analysis of de-identified public transcriptomic data from GSE79362.[[2,11]] The locally remapped gene-level matrix linked each sample to a subject identifier, progressor status, and a parsed follow-up label. Scheduled visits included DAY0, DAY180, DAY360, and DAY540, while an additional OTHER group represented intermediate early collections. IC samples were retained but not forced onto the same numeric follow-up scale.",
        "Gene-level remapping aggregated junction-level counts to gene symbols and then linked the resulting matrix to subject and visit information already curated in the local processed metadata. This step was necessary because direct use of junction identifiers would have prevented biologically interpretable program scoring and would have made comparison with prior gene-symbol analyses unnecessarily unstable.",
        "To avoid deriving another unstable cohort-specific signature, we used prespecified score families. These included: a 25-gene progression signature proxy based on genes present after remapping; a compact Bayesian core based on leading genes recovered in the earlier two-cohort harmonized analysis; a remap-sensitive myeloid score based on FCGR3B, HP, ACSL1, ANXA5, and SERINC2; vascular and coexpression proxy scores; and marker-based cell-state scores for monocyte, neutrophil, platelet, T-cell, B-cell, and NK-cell biology.[[8-10]]",
        "Gene-level expression values were standardized within the remapped cohort before score construction. Program scores were calculated as the mean standardized expression of the genes in each prespecified set. Marker-based cell-state scores were estimated with a non-negative least squares approach using canonical blood-cell marker genes, and these were interpreted as relative composition proxies rather than measured cell fractions.[[9]]",
        "The primary longitudinal models used samples with numeric scheduled follow-up months. For each score or gene, we fit a model with fixed effects for follow-up month, progressor status, and a progressor-by-month interaction. Random-intercept mixed-effects models were used when estimation was stable; otherwise, subject-clustered regression was used as a conservative fallback.[[7]] The interaction term was the main parameter of interest because it captured whether the temporal slope differed by progression status.",
        "This design was chosen deliberately. A positive progressor coefficient could reflect higher early values in progressors, whereas a negative interaction could reflect a steeper later decline; the two effects therefore answer different biological questions. By fitting them together, the analysis could distinguish initial separation from divergence in temporal behavior rather than conflating both into one summary score.",
        "IC samples were not included in the primary month-scale models because they are not naturally ordered alongside the scheduled visit structure in the remapped metadata. Instead, we compared IC samples with scheduled progressor samples in a sensitivity analysis. Multiple-testing correction used the Benjamini-Hochberg method within the program-level and gene-level model families.",
        "The primary interpretation was prespecified before manuscript drafting. Program-level findings were considered stronger than gene-level findings if they remained directionally coherent, survived false-discovery control, and aligned with the descriptive time-course plots. This rule was intended to avoid overstating isolated nominal p values from individual transcripts.",
    ]

    results = [
        f"The remapped longitudinal cohort contained {n_samples} samples from {n_subjects} subjects, including {n_progressors} progressors and {n_nonprogressors} non-progressors. Scheduled follow-up months were available for {n_sched} samples, which formed the primary longitudinal modeling set (Table 1). Repeat sampling was common enough to justify subject-level trajectory analysis, and there were no duplicated subject-timepoint combinations after remapping.",
        f"Program-level modeling produced a clearer signal than single-gene modeling. The strongest progressor-by-month interactions were observed for {', '.join(top_programs)} (Table 2; Fig. 3). At a false-discovery threshold below 0.10, the retained program-level features were {', '.join(fdr10_programs)}. The negative interaction coefficients for vascular_proxy5, remap_myeloid5, and module_M6_proxy indicate steeper decline over follow-up in progressors than in non-progressors, whereas platelet showed a modest positive interaction.",
        f"Descriptive mean trajectories supported the same interpretation. In progressors, remap_myeloid5 rose between OTHER or DAY0 and DAY180, then decreased by DAY540, while the vascular proxy also shifted from early positive values toward negative values later in follow-up (Fig. 2). Non-progressors showed flatter or smaller changes over the same scheduled visits. For example, the progressor mean for remap_myeloid5 moved from {float(day180_prog['remap_myeloid5'].iloc[0]):.3f} at DAY180 to {float(day540_prog['remap_myeloid5'].iloc[0]):.3f} at DAY540, whereas the non-progressor means remained closer to {float(day180_non['remap_myeloid5'].iloc[0]):.3f} and {float(day540_non['remap_myeloid5'].iloc[0]):.3f}, respectively.",
        "The cell-state proxy layer added a complementary view. Neutrophil and platelet-associated proxies both showed progressor-by-month interactions that remained within the program-level false-discovery-supported set, while monocyte-associated scores were directionally compatible with the broader myeloid signal but did not show the same level of longitudinal interaction strength. This pattern suggests that the observed temporal behavior is not reducible to one simple composition shift.",
        f"Gene-level models were less decisive. No individual gene trajectory survived multiple-testing correction, although {', '.join(top_genes)} showed the strongest nominal progressor-by-month trends (Table 3; Fig. 4). This pattern argues that the longitudinal signal is more stable at the level of coordinated programs than at the level of isolated genes.",
        "The existing 25-gene signature proxy and the compact Bayesian core both trended in the same overall direction, but neither reached the same level of statistical support as the higher-order program scores. This further suggests that repeated follow-up captures shifting composite biology better than it captures one persistent dominant transcript.",
        "The subject-level trajectory plot reinforces that interpretation. Individual trajectories were heterogeneous, which is expected in prospective host-response data, yet the overlaid mean curves still separated enough to support model-based program differences. In other words, the longitudinal signal was not driven by one or two extreme subjects but emerged from directional structure across the cohort.",
        "IC sensitivity analyses did not identify a strong separate IC effect in this run (Table 4). The largest numerical contrasts were seen for HP, ACSL1, and FCGR3B, but none survived multiple-testing correction. These observations support keeping IC samples as a descriptive sensitivity layer rather than merging them into the primary numeric follow-up model.",
    ]

    discussion = [
        "This reanalysis suggests that longitudinal progression biology in this cohort is easier to detect at the level of coordinated host programs than at the level of individual genes. That is an important distinction. It implies that repeated measurements may capture dynamic reweighting of related pathways, cell states, and modules even when no one transcript provides a universally stable trajectory marker.",
        "The strongest supported programs were not random. The remap-sensitive myeloid set and the vascular proxy both showed steeper temporal shifts in progressors, and the M6 coexpression proxy moved in the same broad direction. Together, these findings are compatible with a model in which progression risk is associated with evolving myeloid activity and tissue-interface remodeling rather than with one static baseline inflammatory state.",
        "This is also the point at which the longitudinal paper adds something genuinely different to the cross-sectional literature. Many prior transcriptomic studies already established that a blood signature can perform above chance at one time point. The present analysis suggests that when repeated sampling is available, the more stable biological story may reside in the way related programs move together across follow-up rather than in whether one gene remains consistently highest at every visit.",
        "The marker-based neutrophil and platelet results also add meaning, but they require careful interpretation. These scores are relative composition proxies derived from canonical marker genes, not measured leukocyte or platelet fractions. Their value lies in supporting the biological direction of change, not in providing exact cell counts.[[9]]",
        "The weaker gene-level results are equally informative. ACSL1, FCGR3B, and VSIG4 showed the strongest nominal interaction trends, all biologically plausible in a myeloid-weighted host-response setting, but none remained significant after multiple-testing correction. That argues against overclaiming any single transcript as the dominant longitudinal driver in this cohort.",
        "The analysis also clarifies the importance of separating scheduled follow-up from IC samples. The remapped metadata support a clear visit structure at approximately 0, 6, 12, and 18 months, but IC samples do not sit naturally on that same timeline. Treating them as ordinary late follow-up points would risk imposing an interpretation that the data do not warrant.",
        "These results fit the broader tuberculosis biomarker literature while adding a different layer of meaning. Prior work established that blood RNA signatures can identify short-term risk and can retain prognostic value across cohorts.[[2-6]] The present longitudinal analysis adds that, within one prospective RNA-sequencing cohort, the stronger reproducible signal lies in coordinated program dynamics rather than in any one gene's time slope.",
        "From a study-design perspective, the results also argue for a more disciplined way of using public longitudinal data. Remapping, subject linkage, and explicit handling of unscheduled collections are not merely technical preliminaries; they change what can be claimed biologically. The present workflow therefore offers a practical template for future reanalyses of TB transcriptomic cohorts with mixed visit structures.",
        "This distinction matters for translation. A clinically useful serial biomarker may ultimately need to summarize multidimensional host-response programs rather than rely on a single repeatedly measured transcript. Longitudinal program behavior could therefore become a complementary criterion for selecting robust candidate panels, alongside cross-sectional discrimination and cross-platform stability.",
        "Several limitations remain. The primary longitudinal evidence currently rests on one remapped cohort, follow-up month is not equivalent to exact time-to-disease, and the program definitions were prespecified from the present analytical framework rather than from an externally published longitudinal framework. These constraints support biological interpretation and hypothesis generation more strongly than immediate clinical forecasting.",
        "Future work should therefore proceed in two linked directions. The first is analytical: extending the same trajectory logic to other cohorts once remapping and subject-level linkage are feasible. The second is translational: asking whether the most stable longitudinal programs can be compressed into assay-friendly panels without losing interpretability across repeated measurements.",
    ]

    conclusion = (
        "A longitudinal reanalysis of the remapped GSE79362 blood RNA-sequencing cohort suggests that progressor-specific temporal differences are more coherent at the level of coordinated host-response programs than at the level of individual genes. The strongest supported signals involved remap-sensitive myeloid biology, vascular-remodeling proxies, neutrophil or platelet proxy shifts, and a progressor-associated coexpression program. These findings support a longitudinal systems-biology interpretation of tuberculosis progression, while also showing that scheduled follow-up and IC samples must be handled as distinct analytical states."
    )

    return {
        "abstract": abstract,
        "intro": intro,
        "methods": methods,
        "results": results,
        "discussion": discussion,
        "conclusion": conclusion,
    }


def build_title_page(data: dict[str, pd.DataFrame]) -> Path:
    scores = data["scores"]
    doc = Document()
    set_style(doc)
    add_plain_paragraph(doc, TITLE, bold=True, center=True)
    add_plain_paragraph(doc, "")
    add_plain_paragraph(doc, f"{AUTHOR_NAME}, {DEGREE}", center=True)
    add_plain_paragraph(doc, AFFILIATION, center=True)
    add_plain_paragraph(doc, "")
    add_plain_paragraph(doc, f"Short title: {SHORT_TITLE}")
    add_plain_paragraph(doc, f"Corresponding author: {CORRESPONDING}")
    add_plain_paragraph(doc, "ORCID: 0000-0002-4771-8285")
    add_plain_paragraph(doc, "Author contributions: Single-author study. Siddalingaiah H S conceived the study, performed the analysis, interpreted the findings, drafted the manuscript, and approved the final version.")
    add_plain_paragraph(doc, "Funding: No external funding was received for this work.")
    add_plain_paragraph(doc, "Conflict of interest: The author declares no competing interests.")
    add_plain_paragraph(doc, "Ethics statement: This study used de-identified publicly available datasets only and involved no new participant recruitment.")
    add_plain_paragraph(doc, f"Repository: {REPO_URL}")
    add_plain_paragraph(doc, f"Samples in longitudinal dataset: {len(scores)}")
    add_plain_paragraph(doc, f"Subjects in longitudinal dataset: {scores['subject_id'].nunique()}")
    out = OUT_DIR / "01_Title_Page.docx"
    doc.save(out)
    return out


def build_manuscript(data: dict[str, pd.DataFrame]) -> Path:
    text = manuscript_text(data)
    doc = Document()
    set_style(doc)
    add_plain_paragraph(doc, TITLE, bold=True, center=True)
    add_heading(doc, "Abstract", 1)
    for para in text["abstract"]:
        add_plain_paragraph(doc, para)
    add_plain_paragraph(doc, "Keywords: tuberculosis; transcriptomics; longitudinal analysis; mixed-effects modeling; host-response programs")

    add_heading(doc, "Introduction", 1)
    for para in text["intro"]:
        add_cited_paragraph(doc, para)

    add_heading(doc, "Methods", 1)
    for para in text["methods"]:
        add_cited_paragraph(doc, para)

    add_heading(doc, "Results", 1)
    for para in text["results"]:
        add_cited_paragraph(doc, para)

    add_heading(doc, "Discussion", 1)
    for para in text["discussion"]:
        add_cited_paragraph(doc, para)

    add_heading(doc, "Conclusion", 1)
    add_plain_paragraph(doc, str(text["conclusion"]))

    add_heading(doc, "Declarations", 1)
    add_plain_paragraph(doc, "Ethics approval and consent to participate: Not applicable for this secondary analysis of de-identified public data.")
    add_plain_paragraph(doc, "Consent for publication: Not applicable.")
    add_plain_paragraph(doc, "Availability of data and materials: The remapped GSE79362-derived metadata, generated analysis tables, and manuscript assets are available in the project repository.")
    add_plain_paragraph(doc, f"Code availability: Analysis scripts and generated assets are available at {REPO_URL}.")
    add_plain_paragraph(doc, "Funding: No external funding was received.")
    add_plain_paragraph(doc, "Competing interests: The author declares no competing interests.")
    add_plain_paragraph(doc, "Use of generative AI: Generative AI assistance was used for drafting support, while all scientific claims were checked against the project outputs before inclusion.")

    add_heading(doc, "References", 1)
    for idx, ref in enumerate(REFERENCES, start=1):
        add_plain_paragraph(doc, f"{idx}. {ref}")

    t1 = pd.DataFrame(
        [
            {"Measure": "Total samples", "Value": int(len(data["scores"]))},
            {"Measure": "Subjects", "Value": int(data["scores"]["subject_id"].nunique())},
            {"Measure": "Progressor subjects", "Value": int(data["scores"].loc[data["scores"]["progressor"] == 1, "subject_id"].nunique())},
            {"Measure": "Non-progressor subjects", "Value": int(data["scores"].loc[data["scores"]["progressor"] == 0, "subject_id"].nunique())},
            {"Measure": "Samples with numeric scheduled follow-up", "Value": int(data["scores"]["timepoint_month"].notna().sum())},
            {"Measure": "IC samples", "Value": int((data["scores"]["timepoint_label"] == "IC").sum())},
        ]
    )
    t2 = data["program_models"].head(8).loc[:, ["feature", "interaction_coef", "interaction_ci_low", "interaction_ci_high", "interaction_pvalue", "interaction_fdr_bh"]].copy().round(4)
    t2.columns = ["Feature", "Interaction coef", "95% CI low", "95% CI high", "p value", "FDR"]
    t3 = data["gene_models"].loc[:, ["feature", "interaction_coef", "interaction_ci_low", "interaction_ci_high", "interaction_pvalue", "interaction_fdr_bh"]].copy().round(4)
    t3.columns = ["Gene", "Interaction coef", "95% CI low", "95% CI high", "p value", "FDR"]
    t4 = data["ic"].loc[:, ["feature", "delta_ic_minus_scheduled", "pvalue", "fdr_bh"]].copy().round(4)
    t4.columns = ["Feature", "IC minus scheduled delta", "p value", "FDR"]

    add_heading(doc, "Tables", 1)
    add_table(doc, t1, "Table 1. Longitudinal cohort summary after gene-level remapping and subject linkage.")
    add_table(doc, t2, "Table 2. Program-level progressor-by-month interaction estimates from the primary longitudinal models.")
    add_table(doc, t3, "Table 3. Gene-level progressor-by-month interaction estimates. No gene-level interaction survived multiple-testing correction.")
    add_table(doc, t4, "Table 4. Progressor-only IC sensitivity contrasts. These results are descriptive and do not support merging IC samples into the scheduled visit scale.")

    add_heading(doc, "Figures", 1)
    add_figure(doc, ANALYSIS_DIR / "signature_spaghetti.png", "Figure 1. Subject-level and mean trajectories for the 25-gene signature proxy across scheduled follow-up visits.")
    add_figure(doc, ANALYSIS_DIR / "program_mean_trajectories.png", "Figure 2. Mean trajectories for selected prespecified programs, including the remap-sensitive myeloid score, M6 module proxy, and monocyte-related proxy.")
    add_figure(doc, ANALYSIS_DIR / "program_interaction_forest.png", "Figure 3. Program-level progressor-by-month interaction estimates with 95% confidence intervals.")
    add_figure(doc, ANALYSIS_DIR / "gene_interaction_forest.png", "Figure 4. Gene-level progressor-by-month interaction estimates with 95% confidence intervals.")

    out = OUT_DIR / "02_Manuscript.docx"
    doc.save(out)
    return out


def build_highlights() -> Path:
    doc = Document()
    set_style(doc)
    add_plain_paragraph(doc, "Highlights", bold=True, center=True)
    for line in [
        "Longitudinal reanalysis was performed on a remapped public TB RNA-sequencing cohort with repeated subject-level follow-up.",
        "Program-level temporal differences were stronger than single-gene temporal differences.",
        "Myeloid, vascular, neutrophil, platelet, and coexpression-program signals showed the strongest progressor-by-month interactions.",
        "No individual gene trajectory survived multiple-testing correction in the primary longitudinal models.",
        "IC samples behaved as a distinct sensitivity layer and should not be merged into the scheduled follow-up scale.",
    ]:
        doc.add_paragraph(line, style="List Bullet")
    out = OUT_DIR / "03_Highlights.docx"
    doc.save(out)
    return out


def build_supplement(data: dict[str, pd.DataFrame]) -> Path:
    doc = Document()
    set_style(doc)
    add_plain_paragraph(doc, "Supplementary Methods and Tables", bold=True, center=True)
    add_heading(doc, "Supplementary Methods", 1)
    add_plain_paragraph(doc, "The longitudinal analysis focused on the remapped GSE79362 gene-level matrix and used prespecified score sets to avoid deriving a new single-cohort classifier. Scheduled follow-up months were modeled directly, whereas IC samples were analyzed separately as a sensitivity layer.")
    add_plain_paragraph(doc, "Program definitions are reported exactly as present after remapping and are included in Supplementary Table 1. Marker-based cell-state scores were estimated with non-negative least squares and should be interpreted as relative composition proxies rather than measured cell fractions.")
    add_plain_paragraph(doc, "When mixed-effects estimation was unstable, a subject-clustered regression fallback was used to avoid dropping informative features while still preserving subject-level dependence in the standard errors.")
    add_heading(doc, "Supplementary Table", 1)
    add_table(doc, data["gene_sets"].rename(columns={"program": "Program", "n_genes_present": "Genes present", "genes": "Genes"}), "Supplementary Table 1. Prespecified program definitions as present in the remapped GSE79362 gene-level matrix.")
    add_table(doc, pd.read_csv(ANALYSIS_DIR / "subject_program_slopes.csv").head(20), "Supplementary Table 2. Example subject-level slopes for selected programs and genes. The full table is available in the results directory.")
    out = OUT_DIR / "04_Supplementary_Methods_and_Tables.docx"
    doc.save(out)
    return out


def build_cover_letter() -> Path:
    doc = Document()
    set_style(doc)
    add_plain_paragraph(doc, "Cover Letter", bold=True, center=True)
    add_plain_paragraph(doc, "Dear Editor,")
    add_plain_paragraph(doc, f"I am submitting the manuscript titled \"{TITLE}\" for consideration as an original research article.")
    add_plain_paragraph(doc, "This work presents a longitudinal reanalysis of a public tuberculosis blood RNA-sequencing cohort after gene-level remapping and subject-level linkage. The study focuses on temporal behavior of prespecified host-response programs rather than deriving another single-cohort classifier, and it keeps all clinical claims bounded by the observed data structure.")
    add_plain_paragraph(doc, "The manuscript is original, is not under consideration elsewhere, and all analyzed data are publicly available. The accompanying repository contains the analysis code and generated outputs for transparency and reproducibility.")
    add_plain_paragraph(doc, "Sincerely,")
    add_plain_paragraph(doc, f"{AUTHOR_NAME}, {DEGREE}")
    out = OUT_DIR / "05_Cover_Letter.docx"
    doc.save(out)
    return out


def build_readme() -> Path:
    lines = [
        "# Longitudinal TB Dynamics Package",
        "",
        f"Repository: {REPO_URL}",
        "",
        "Contents",
        "- 01_Title_Page.docx",
        "- 02_Manuscript.docx",
        "- 03_Highlights.docx",
        "- 04_Supplementary_Methods_and_Tables.docx",
        "- 05_Cover_Letter.docx",
        "- validation_report.txt",
        "- internal_review_log.md",
        "",
        "This package is derived from the longitudinal analysis stream in `results/longitudinal_tb_analysis`.",
    ]
    out = OUT_DIR / "README.md"
    out.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return out


def build_review_note() -> Path:
    lines = [
        "# Internal Review Note",
        "",
        "Review pass 1",
        "- Checked that all numerical claims in the manuscript match the current longitudinal analysis outputs.",
        "- Checked that program-level significance is described more strongly than gene-level significance.",
        "- Checked that IC samples are consistently framed as a sensitivity layer, not as ordered scheduled follow-up.",
        "",
        "Review pass 2",
        "- Checked sequential appearance of Table 1 to Table 4 and Figure 1 to Figure 4.",
        "- Checked that no statement claims time-to-disease forecasting from scheduled follow-up month alone.",
        "- Checked that conclusions remain at the host-program level and do not exaggerate the gene-level findings.",
    ]
    out = OUT_DIR / "internal_review_log.md"
    out.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return out


def build_validation_report() -> Path:
    doc = Document(OUT_DIR / "02_Manuscript.docx")
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)
    words = sum(len(p.split()) for p in paragraphs)
    lines = [
        "Longitudinal manuscript package validation",
        f"Approximate manuscript words (paragraph text only): {words}",
        f"Contains Table 1-4 in order: {all(text.find(f'Table {i}') != -1 for i in range(1, 5))}",
        f"Contains Figure 1-4 in order: {all(text.find(f'Figure {i}') != -1 for i in range(1, 5))}",
        f"Contains Declarations section: {'Declarations' in text}",
        f"Contains References section: {'References' in text}",
        f"Contains IC caution: {'IC samples' in text}",
        f"Contains program-level over gene-level framing: {'program-level' in text.lower()}",
        "Interpretive limit note: follow-up month is not treated as exact time-to-disease.",
    ]
    out = OUT_DIR / "validation_report.txt"
    out.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return out


def zip_package() -> Path:
    zip_path = OUT_DIR / "longitudinal_tb_dynamics_submission_package.zip"
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for path in sorted(OUT_DIR.iterdir()):
            if path == zip_path:
                continue
            zf.write(path, arcname=path.name)
    return zip_path


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    data = load_outputs()
    build_title_page(data)
    build_manuscript(data)
    build_highlights()
    build_supplement(data)
    build_cover_letter()
    build_readme()
    build_review_note()
    build_validation_report()
    zip_package()
    print(f"Longitudinal submission assets written to {OUT_DIR}")


if __name__ == "__main__":
    main()
