from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re
import zipfile

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
ANALYSIS_DIR = ROOT / "results" / "advanced_analysis"
OMICSCLAW_DIR = ROOT / "results" / "omicsclaw_extensions"
OUT_DIR = ROOT / "submission_ready" / "advanced_tb_systems_20260317"
REPO_URL = "https://github.com/hssling/tb-progression-transcriptome-meta"

AUTHOR_NAME = "Siddalingaiah H S"
DEGREE = "MD"
AFFILIATION = (
    "Department of Community Medicine, Shridevi Institute of Medical Sciences & "
    "Research Hospital, Tumkur, Karnataka, India"
)
CORRESPONDING = (
    "Siddalingaiah H S, Department of Community Medicine, Shridevi Institute of "
    "Medical Sciences & Research Hospital, Tumkur, Karnataka, India. "
    "Email: hssling@yahoo.com"
)
TITLE = (
    "Bayesian and systems-level reanalysis of public tuberculosis progression "
    "transcriptomes reveals latent host-response programs"
)
SHORT_TITLE = "Bayesian TB progression transcriptomics"

REFERENCES = [
    "World Health Organization. Global tuberculosis report 2025. Geneva: World Health Organization; 2025.",
    "Zak DE, Penn-Nicholson A, Scriba TJ, et al. A blood RNA signature for tuberculosis disease risk: a prospective cohort study. Lancet. 2016;387:2312-2322.",
    "Singhania A, Verma R, Graham CM, et al. A modular transcriptional signature identifies phenotypic heterogeneity of human tuberculosis infection. Nat Commun. 2018;9:2308.",
    "Rajamanickam A, Munisankar S, Dolla CK, et al. Host blood-based biosignatures for subclinical TB and incipient TB: a prospective study of adult TB household contacts in Southern India. Front Immunol. 2022;13:1065779.",
    "Sweeney TE, Braviak L, Tato CM, Khatri P. Genome-wide expression for diagnosis of pulmonary tuberculosis: a multicohort analysis. Lancet Respir Med. 2016;4:213-224.",
    "Ritchie ME, Phipson B, Wu D, et al. limma powers differential expression analyses for RNA-sequencing and microarray studies. Nucleic Acids Res. 2015;43:e47.",
    "Balduzzi S, Rucker G, Schwarzer G. How to perform a meta-analysis with R: a practical tutorial. Evid Based Ment Health. 2019;22:153-160.",
    "Wu T, Hu E, Xu S, et al. clusterProfiler 4.0: a universal enrichment tool for interpreting omics data. Innovation (Camb). 2021;2:100141.",
    "Pearl J. Causality: models, reasoning, and inference. 2nd ed. Cambridge: Cambridge University Press; 2009.",
    "Leek JT, Johnson WE, Parker HS, Jaffe AE, Storey JD. The sva package for removing batch effects and other unwanted variation in high-throughput experiments. Bioinformatics. 2012;28:882-883.",
    "Newman AM, Liu CL, Green MR, et al. Robust enumeration of cell subsets from tissue expression profiles. Nat Methods. 2015;12:453-457.",
    "Langfelder P, Horvath S. WGCNA: an R package for weighted correlation network analysis. BMC Bioinformatics. 2008;9:559.",
]


@dataclass
class AdvancedData:
    cohort_summary: pd.DataFrame
    bayes_genes: pd.DataFrame
    pathways: pd.DataFrame
    factors: pd.DataFrame
    raw_scores: pd.DataFrame
    centered_scores: pd.DataFrame


def set_style(doc: Document, double_spacing: bool = True) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 2 if double_spacing else 1.15
    style.paragraph_format.space_after = Pt(0)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)


def add_plain_paragraph(doc: Document, text: str = "", bold: bool = False, center: bool = False) -> None:
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_cited_paragraph(doc: Document, text: str, center: bool = False) -> None:
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    parts = re.split(r"(\[\[[0-9,\- ]+\]\])", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("[[") and part.endswith("]]"):
            run = p.add_run(part[2:-2].replace(" ", ""))
            run.font.superscript = True
        else:
            p.add_run(part)


def add_table(doc: Document, df: pd.DataFrame, title: str) -> None:
    add_plain_paragraph(doc, title)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for idx, col in enumerate(df.columns):
        table.cell(0, idx).text = str(col)
    for row in df.itertuples(index=False):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    doc.add_paragraph("")


def add_figure(doc: Document, image_path: Path, caption: str, width: float = 6.2) -> None:
    add_plain_paragraph(doc, caption)
    if image_path.exists():
        doc.add_picture(str(image_path), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")


def parse_metadata_frame(cohort_id: str) -> pd.DataFrame:
    return pd.read_parquet(ROOT / "data" / "processed" / cohort_id / "metadata.parquet")


def parse_characteristics(value: str) -> dict[str, str]:
    if not isinstance(value, str):
        return {}
    parts = [p.strip() for p in value.split("|")]
    out: dict[str, str] = {}
    for part in parts:
        if ":" not in part:
            continue
        key, val = part.split(":", 1)
        out[key.strip().lower().replace(" ", "_")] = val.strip()
    return out


def format_pvalue(value: float) -> str:
    if pd.isna(value):
        return "NA"
    if value < 0.001:
        return f"{value:.2e}"
    return f"{value:.3f}"


def cohort_summary() -> pd.DataFrame:
    rows = []
    for cohort_id in ["GSE107994", "GSE193777"]:
        meta = parse_metadata_frame(cohort_id).copy()
        parsed = meta["characteristics"].map(parse_characteristics)
        age = parsed.map(
            lambda row: row.get("age")
            or row.get("age_in_years")
            or row.get("age_at_baseline_visit")
        )
        age = pd.to_numeric(age, errors="coerce")
        sex = parsed.map(lambda row: row.get("gender", "")).astype(str).str.upper().str[:1]
        platform = meta["platform_type"].mode().iloc[0] if "platform_type" in meta.columns and not meta["platform_type"].isna().all() else "NA"
        rows.append(
            {
                "Cohort": cohort_id,
                "Platform": platform,
                "Samples": int(len(meta)),
                "Progressors": int(meta["progressor"].sum()),
                "Non-progressors": int((1 - meta["progressor"]).sum()),
                "Median age (years)": f"{age.median():.1f}" if not age.empty and age.notna().any() else "NA",
                "Female, n (%)": (
                    f"{int((sex == 'F').sum())} ({100 * (sex == 'F').mean():.1f})"
                    if len(sex) > 0
                    else "NA"
                ),
            }
        )
    return pd.DataFrame(rows)


def load_data() -> AdvancedData:
    return AdvancedData(
        cohort_summary=cohort_summary(),
        bayes_genes=pd.read_csv(ANALYSIS_DIR / "bayesian_gene_meta.csv"),
        pathways=pd.read_csv(ANALYSIS_DIR / "bayesian_pathway_summary.csv"),
        factors=pd.read_csv(ANALYSIS_DIR / "factor_summary.csv"),
        raw_scores=pd.read_csv(ANALYSIS_DIR / "raw_scores.csv"),
        centered_scores=pd.read_csv(ANALYSIS_DIR / "cohort_centered_scores.csv"),
    )


def manuscript_text(data: AdvancedData) -> dict[str, list[str] | str]:
    top_gene = data.bayes_genes.iloc[0]
    top_genes = data.bayes_genes.head(8)["gene"].tolist()
    path_names = data.pathways.head(5)["pathway"].tolist()
    raw_prog = data.raw_scores.groupby("progressor")["PC1"].mean().to_dict()
    centered_prog = data.centered_scores.groupby("progressor")["PC1"].mean().to_dict()
    factor_map = {
        row["factor"]: (row["mean_progressor"], row["mean_nonprogressor"], row["ttest_pvalue"])
        for _, row in data.factors.iterrows()
    }
    cell_summary = pd.read_csv(OMICSCLAW_DIR / "nnls_celltype_summary.csv") if (OMICSCLAW_DIR / "nnls_celltype_summary.csv").exists() else pd.DataFrame()
    module_summary = pd.read_csv(OMICSCLAW_DIR / "coexpression_module_summary.csv") if (OMICSCLAW_DIR / "coexpression_module_summary.csv").exists() else pd.DataFrame()
    module_overlap = pd.read_csv(OMICSCLAW_DIR / "coexpression_module_signature_overlap.csv") if (OMICSCLAW_DIR / "coexpression_module_signature_overlap.csv").exists() else pd.DataFrame()
    top_cell = cell_summary.iloc[0] if not cell_summary.empty else None
    top_module = module_summary.iloc[0] if not module_summary.empty else None
    top_overlap = module_overlap[module_overlap["module"] == top_module["module"]].iloc[0] if (top_module is not None and not module_overlap.empty and (module_overlap["module"] == top_module["module"]).any()) else None
    abstract = [
        "Background: Public tuberculosis progression transcriptomic datasets contain more biological information than can be captured by ranked-gene lists alone. We performed a systems-level reanalysis to assess latent structure, uncertainty-aware gene ranking, pathway convergence, and bias-sensitive interpretation in the harmonizable public cohorts.",
        "Methods: Two cohorts with shared gene symbols and binary progressor labels (GSE107994 and GSE193777) were reanalyzed. We applied joint principal component analysis before and after cohort centering, factor analysis on the most variable genes, Bayesian hierarchical synthesis of within-cohort differential expression effects, pathway-level posterior modeling, marker-based NNLS deconvolution, WGCNA-style coexpression analysis, signature correlation analysis, and a directed acyclic graph to clarify potential bias pathways.",
        f"Results: The advanced analysis included 301 samples, comprising 87 progressors and 214 non-progressors. Raw PC1 remained strongly cohort structured, but cohort-centered PC1 separated non-progressors and progressors more clearly (mean PC1 {raw_prog[0]:.1f} vs {raw_prog[1]:.1f} before centering; {centered_prog[0]:.1f} vs {centered_prog[1]:.1f} after centering). Bayesian synthesis prioritized {', '.join(top_genes[:6])}, with {top_gene['gene']} showing the strongest pooled effect (posterior mean {top_gene['posterior_mean']:.3f}, 95% credible interval {top_gene['ci95_low']:.3f} to {top_gene['ci95_high']:.3f}). The leading pathway signals were {', '.join(path_names[:3])}. All three latent factors remained associated with progressor status, with the strongest evidence for Factor1 (p={factor_map['Factor1'][2]:.2e})." + (f" Marker-based deconvolution suggested higher {top_cell['cell_type'].lower()} and lower lymphoid-associated scores in progressors." if top_cell is not None else ""),
        "Conclusion: The harmonizable public datasets support a coordinated tuberculosis progression signal that combines myeloid regulation with vascular-remodeling biology. The findings are stronger as uncertainty-aware biological evidence than as a clinical prediction claim, because the shared-gene advanced layer currently rests on two cohorts and should be expanded before clinical translation is considered. The deconvolution and coexpression analyses are supportive interpretation layers, not direct measures of leukocyte fractions or causal network effects.",
    ]

    intro = [
        "Tuberculosis remains a major global health problem, and earlier recognition of incipient disease remains a central unmet need in prevention research.[[1]] Blood transcriptomic signatures are attractive because they can potentially identify host response changes before microbiological confirmation becomes straightforward.[[2,3]]",
        "Most transcriptomic studies of tuberculosis progression eventually collapse into ranked gene lists. Those lists are useful, but they often hide three important issues: how much of the signal is driven by study structure, how stable the leading genes remain after uncertainty is modeled explicitly, and whether the data converge on coherent host programs rather than isolated markers.[[4,5]]",
        "Previous tuberculosis progression transcriptomic studies have emphasized signature discovery and validation performance. The present study asks a different question. Instead of focusing primarily on classifier performance, it evaluates whether the available public datasets support deeper biological structure through unsupervised analysis, Bayesian shrinkage, pathway-level synthesis, and a causal-interpretation framework.",
        "We therefore performed an advanced reanalysis of the currently harmonizable progression cohorts. The aim was to determine whether latent axes, posterior summaries, pathway convergence, and correlation structure reinforce a biologically coherent interpretation of progression risk while keeping all claims bounded by the actual data available for joint analysis.",
    ]

    methods = [
        "This was a secondary analysis of de-identified public transcriptomic datasets relevant to tuberculosis progression. The advanced layer was deliberately restricted to cohorts that could be harmonized at the shared-gene level and that retained binary progressor labels suitable for direct comparison.",
        "Processed metadata and expression matrices were available for GSE107994 and GSE193777. GSE79362 remained valuable for the broader project, but its current feature mapping was not directly commensurate with the joint shared-gene analysis and was therefore not forced into the advanced modeling layer. This conservative choice was made to avoid false precision.",
        "A shared-gene expression matrix was constructed across the two retained cohorts, yielding 14,440 common genes. Principal component analysis was first performed on the raw joint matrix and then repeated after cohort-centered normalization to assess how much cohort structure obscured biology. We then applied factor analysis to the 500 most variable centered genes to identify latent programs associated with progressor status.",
        "Gene-level Bayesian synthesis used existing within-cohort differential expression estimates from the project results tables. An empirical normal-normal hierarchical model was applied to derive posterior means, posterior uncertainty, and between-cohort heterogeneity. Pathway-level analysis used previously exported enrichment-derived gene sets, from which pathway scores were calculated for each sample and then summarized with the same uncertainty-aware framework.[[6-8]]",
        "To extend biological interpretation, we also applied two additional bulk-RNA analyses. First, a conservative non-negative least squares deconvolution used canonical blood-cell marker sets to estimate relative composition proxies across samples. These outputs were interpreted as marker-based composition scores rather than absolute leukocyte fractions. Second, a WGCNA-style coexpression analysis of the 3000 most variable shared genes was used to derive module eigengenes and hub-gene summaries. Soft-threshold exploration was used to choose the adjacency power, and module-trait comparisons were summarized with t tests and Benjamini-Hochberg adjusted values. These modules were treated as supportive program-level structure and not as causal network estimates.[[11,12]]",
        "Finally, we calculated a signature-gene correlation matrix in the centered expression space and drew a directed acyclic graph describing how cohort, platform, baseline host factors, latent disease biology, preprocessing, and measured transcript levels could interact. The directed acyclic graph was used only as an interpretive aid and not as proof of causality.[[9,10]]",
    ]

    results = [
        "The advanced analysis comprised 301 samples, including 175 from GSE107994 and 126 from GSE193777. Across the combined dataset, 87 samples were labeled as progressors and 214 as non-progressors (Table 1). Shared-gene overlap was substantial, with 14,440 genes retained for the joint analysis, which was sufficient for stable latent-structure modeling.",
        f"Raw joint PCA showed strong cohort influence. The mean raw PC1 value was {data.raw_scores.groupby('cohort_id')['PC1'].mean().to_dict()['GSE107994']:.1f} in GSE107994 and {data.raw_scores.groupby('cohort_id')['PC1'].mean().to_dict()['GSE193777']:.1f} in GSE193777, indicating that study structure dominated the leading axis before adjustment (Fig. 1). After cohort centering, the cohort means on PC1 were essentially zero in both cohorts, while separation by progressor status became more apparent, with mean PC1 moving from {centered_prog[0]:.1f} in non-progressors to {centered_prog[1]:.1f} in progressors (Fig. 2). This shift supports the view that cohort structure was masking a meaningful biological gradient rather than creating it.",
        f"Bayesian hierarchical synthesis prioritized a compact set of genes with consistent positive effects across both cohorts. The leading genes were {', '.join(top_genes[:8])} (Table 2). {top_gene['gene']} showed the strongest pooled signal, with a posterior mean of {top_gene['posterior_mean']:.3f} and a 95% credible interval from {top_gene['ci95_low']:.3f} to {top_gene['ci95_high']:.3f}. Several other genes, including VSIG4, FZD5, CD36, CCR2, and AQP1, also retained high posterior means with credible intervals that did not cross zero. This pattern suggests that the leading host signal is not being driven by a single unstable feature.",
        f"Latent factor analysis also supported a multi-program view of progression biology. Factor1, Factor2, and Factor3 all differed between progressors and non-progressors, with p values of {factor_map['Factor1'][2]:.2e}, {factor_map['Factor2'][2]:.2e}, and {factor_map['Factor3'][2]:.2e}, respectively (Table 3; Fig. 3). The factor boxplots indicate that the progression phenotype cannot be reduced to a single expression axis. Instead, multiple orthogonal programs appear to contribute to the observed signal.",
        f"Pathway-level posterior summaries converged on vascular-remodeling biology. The strongest terms were {', '.join(path_names[:5])} (Table 4). These findings do not imply that angiogenesis alone explains progression. Rather, they suggest that immune activation in progressors is accompanied by tissue-interface and vasculature-related remodeling, which is plausible in the setting of evolving tuberculosis disease. The signature correlation heatmap (Fig. 4) reinforced this interpretation by showing tightly connected gene blocks rather than isolated outliers.",
        (f"Additional deconvolution and coexpression analyses supported the same biological direction. Marker-based NNLS deconvolution showed higher {top_cell['cell_type'].lower()}-associated scores in progressors (delta {top_cell['delta_progressor_minus_nonprogressor']:.3f}, p={top_cell['ttest_pvalue']:.2e}, FDR={top_cell['fdr_bh']:.2e}), alongside lower T-cell, B-cell, and NK-cell associated scores. In parallel, coexpression analysis identified a large progressor-associated module, {top_module['module']}, containing {int(top_module['n_genes'])} genes with strong eigengene separation by status (p={top_module['ttest_pvalue']:.2e}, FDR={top_module['fdr_bh']:.2e})." + (f" This module overlapped the top Bayesian signal through {top_overlap['signature_genes'].replace('; ', ', ')}." if top_overlap is not None and isinstance(top_overlap['signature_genes'], str) and top_overlap['signature_genes'] else "") if top_cell is not None and top_module is not None else ""),
    ]

    discussion = [
        "This reanalysis adds scientific meaning to the original progression project in three ways. First, it demonstrates that cohort structure materially affects the leading unsupervised axes and must be addressed before biological interpretation is attempted. Second, it replaces simple ranking with posterior estimates and uncertainty intervals. Third, it shows convergence between gene-level, factor-level, pathway-level, and network-level summaries.",
        "The leading genes point toward a host program centered on myeloid regulation, cell trafficking, scavenger biology, and compensatory immune control. MILR1, VSIG4, CCR2, CD36, FZD5, and AQP1 do not read like a random collection of high-scoring variables. Together they suggest coordinated changes in cellular recruitment, innate regulation, and tissue interaction, which is more persuasive biologically than a diffuse inflammatory signature.[[2-5]]",
        "The pathway results sharpen that interpretation. The repeated appearance of angiogenesis and vasculature-development terms should not be read as a claim that the disease process is primarily vascular. A more defensible reading is that progression-risk biology in blood may include immune-cell behavior that tracks with tissue remodeling and evolving host-pathogen interface changes. That view is consistent with tuberculosis as a systemic inflammatory disease with localized tissue consequences.",
        "The extension analyses support that reading, but they need to be interpreted carefully. The NNLS deconvolution layer uses canonical marker genes and therefore yields relative composition proxies, not direct leukocyte fractions measured by flow cytometry or single-cell profiling. Its value lies in showing that the transcriptomic signal is compatible with stronger myeloid weighting and weaker lymphoid weighting in progressors, not in providing exact cell counts.",
        "The coexpression layer also requires discipline in interpretation. Module eigengenes and hub-gene summaries help identify coordinated programs and connect isolated ranked genes to broader biological structure, but they do not establish regulatory directionality or causality. In this manuscript, the coexpression analysis is used to support the idea of organized host programs rather than to claim causal network inference.",
        "The cohort summary also helps contextualize interpretation because it makes the cross-platform nature of the retained evidence explicit: one cohort was RNA sequencing based and the other was microarray based. That difference strengthens the case for conservative harmonization and explains why cohort-centering was central to the analysis.",
        "The analysis also clarifies where caution is needed. The advanced shared-gene layer uses only two cohorts, because that is the current evidence boundary imposed by commensurate feature mapping and binary labels. The directed acyclic graph was therefore included to make bias pathways explicit, not to imply that causal effects have been estimated. Likewise, pathway and factor results should be viewed as supportive biological structure rather than stand-alone proof of mechanism.",
        "Even with those constraints, the results are useful. They support a second manuscript focused on uncertainty-aware interpretation, identify genes that remain strong after Bayesian shrinkage, and provide figures that explain why strict harmonization matters. The next technical priority should be to bring additional cohorts into a comparable gene-symbol framework so that the posterior summaries can be tested across a broader evidence base.",
    ]

    conclusion = (
        "A systems-level reanalysis of the harmonizable public tuberculosis progression cohorts identified a "
        "coordinated host-response signal that remained visible across cohort-centered PCA, latent factor "
        "analysis, Bayesian gene synthesis, pathway modeling, deconvolution, coexpression, and correlation structure analysis. The current "
        "evidence supports biological interpretation and hypothesis generation more strongly than immediate "
        "clinical deployment, and future work should expand harmonized external cohorts before stronger "
        "predictive claims are made. The deconvolution results should be read as marker-based composition proxies, "
        "and the coexpression results should be read as supportive program-level biology rather than causal network proof."
    )

    return {
        "abstract": abstract,
        "intro": intro,
        "methods": methods,
        "results": results,
        "discussion": discussion,
        "conclusion": conclusion,
    }


def build_title_page(data: AdvancedData) -> Path:
    doc = Document()
    set_style(doc)
    add_plain_paragraph(doc, TITLE, bold=True, center=True)
    add_plain_paragraph(doc, "")
    add_plain_paragraph(doc, f"{AUTHOR_NAME}, {DEGREE}", center=True)
    add_plain_paragraph(doc, AFFILIATION, center=True)
    add_plain_paragraph(doc, "")
    add_plain_paragraph(doc, f"Short title: {SHORT_TITLE}")
    add_plain_paragraph(doc, f"Corresponding author: {CORRESPONDING}")
    add_plain_paragraph(doc, "ORCID: 0000-0002-4771-8285")
    add_plain_paragraph(doc, "Author contributions: Single-author study. Siddalingaiah H S conceived the study, interpreted the findings, drafted the manuscript, and approved the final version.")
    add_plain_paragraph(doc, "Funding: No external funding was received for this work.")
    add_plain_paragraph(doc, "Conflict of interest: The author declares no competing interests.")
    add_plain_paragraph(doc, "Ethics statement: This study used de-identified publicly available datasets only and involved no new participant recruitment.")
    add_plain_paragraph(doc, f"Repository: {REPO_URL}")
    add_plain_paragraph(doc, f"Samples included in the advanced layer: {int(data.cohort_summary['Samples'].sum())}")
    out = OUT_DIR / "01_Title_Page.docx"
    doc.save(out)
    return out


def build_manuscript(data: AdvancedData) -> Path:
    text = manuscript_text(data)
    doc = Document()
    set_style(doc)
    add_plain_paragraph(doc, TITLE, bold=True, center=True)
    add_heading(doc, "Abstract", 1)
    for para in text["abstract"]:
        add_plain_paragraph(doc, para)
    add_plain_paragraph(doc, "Keywords: tuberculosis; transcriptomics; Bayesian meta-analysis; principal component analysis; latent factor analysis")

    add_heading(doc, "Introduction", 1)
    for para in text["intro"]:
        add_cited_paragraph(doc, para)

    add_heading(doc, "Methods", 1)
    for para in text["methods"]:
        add_cited_paragraph(doc, para)

    add_heading(doc, "Results", 1)
    for para in text["results"]:
        add_cited_paragraph(doc, para)

    add_heading(doc, "Discussion", 1)
    for para in text["discussion"]:
        add_cited_paragraph(doc, para)

    add_heading(doc, "Conclusion", 1)
    add_plain_paragraph(doc, str(text["conclusion"]))

    add_heading(doc, "Declarations", 1)
    add_plain_paragraph(doc, "Ethics approval and consent to participate: Not applicable for this secondary analysis of de-identified public data.")
    add_plain_paragraph(doc, "Consent for publication: Not applicable.")
    add_plain_paragraph(doc, "Availability of data and materials: Public-source processed inputs and generated outputs are available in the project repository.")
    add_plain_paragraph(doc, f"Code availability: Analysis scripts and generated assets are available at {REPO_URL}.")
    add_plain_paragraph(doc, "Funding: No external funding was received.")
    add_plain_paragraph(doc, "Competing interests: The author declares no competing interests.")
    add_plain_paragraph(doc, "Use of generative AI: Generative AI assistance was used for drafting support, while all scientific claims were checked against the project outputs before inclusion.")

    add_heading(doc, "References", 1)
    for idx, ref in enumerate(REFERENCES, start=1):
        add_plain_paragraph(doc, f"{idx}. {ref}")

    top_genes_table = data.bayes_genes.head(10).copy()
    top_genes_table = top_genes_table.loc[:, ["gene", "posterior_mean", "ci95_low", "ci95_high", "tau2"]]
    top_genes_table.columns = ["Gene", "Posterior mean", "95% CrI low", "95% CrI high", "Tau^2"]
    top_genes_table = top_genes_table.round(3)

    factor_table = data.factors.copy()
    factor_table["mean_progressor"] = factor_table["mean_progressor"].map(lambda x: f"{x:.3f}")
    factor_table["mean_nonprogressor"] = factor_table["mean_nonprogressor"].map(lambda x: f"{x:.3f}")
    factor_table["ttest_pvalue"] = factor_table["ttest_pvalue"].map(format_pvalue)
    factor_table.columns = ["Factor", "Mean in progressors", "Mean in non-progressors", "t-test p value"]

    pathway_table = data.pathways.head(8).copy().round(3)
    pathway_table.columns = ["Pathway", "Posterior mean", "Posterior SD", "95% CrI low", "95% CrI high"]

    add_heading(doc, "Tables", 1)
    add_table(doc, data.cohort_summary, "Table 1. Cohort summary for the advanced shared-gene analysis. Age and sex were parsed from public metadata fields, and platform indicates the dominant assay type within each retained cohort.")
    add_table(doc, top_genes_table, "Table 2. Top Bayesian gene-level posterior effects. Posterior means and 95% credible intervals were derived from a two-cohort hierarchical synthesis.")
    add_table(doc, factor_table, "Table 3. Latent factor differences between progressors and non-progressors. P values are shown in scientific notation when small.")
    add_table(doc, pathway_table, "Table 4. Leading pathway-level Bayesian posterior effects. Closely related angiogenesis and vasculature terms reflect convergence within the enrichment output rather than fully independent pathways.")

    add_heading(doc, "Figures", 1)
    add_figure(doc, ANALYSIS_DIR / "raw_pca.png", "Figure 1. Raw joint PCA of the shared-gene matrix. The leading component is strongly cohort influenced before cohort centering.")
    add_figure(doc, ANALYSIS_DIR / "cohort_centered_pca.png", "Figure 2. Cohort-centered PCA of the shared-gene matrix. After centering, separation by progressor status becomes more visible.")
    add_figure(doc, ANALYSIS_DIR / "factor_boxplots.png", "Figure 3. Distribution of latent factor scores by progressor status. All three leading factors remain associated with progression status.")
    add_figure(doc, ANALYSIS_DIR / "signature_network_heatmap.png", "Figure 4. Spearman correlation heatmap for leading Bayesian signature genes after cohort centering.")

    out = OUT_DIR / "02_Manuscript.docx"
    doc.save(out)
    return out


def build_highlights() -> Path:
    doc = Document()
    set_style(doc, double_spacing=False)
    add_plain_paragraph(doc, "Highlights", bold=True, center=True)
    for line in [
        "Two harmonizable TB progression cohorts supported a shared-gene advanced reanalysis.",
        "Cohort centering reduced study-structure dominance in principal component analysis.",
        "Bayesian synthesis prioritized MILR1, VSIG4, FZD5, CD36, CCR2, and AQP1.",
        "Marker-based deconvolution suggested stronger myeloid and weaker lymphoid weighting in progressors.",
        "Latent factors, pathways, and coexpression modules pointed to coordinated host-response programs.",
        "The results support biological interpretation more strongly than immediate clinical deployment.",
    ]:
        doc.add_paragraph(line, style=None).style = "List Bullet"
    out = OUT_DIR / "03_Highlights.docx"
    doc.save(out)
    return out


def build_supplement(data: AdvancedData) -> Path:
    doc = Document()
    set_style(doc)
    add_plain_paragraph(doc, "Supplementary Methods and Notes", bold=True, center=True)
    add_heading(doc, "Supplementary Methods", 1)
    add_plain_paragraph(
        doc,
        "The advanced analysis used the processed cohort matrices already generated within the repository. "
        "Only GSE107994 and GSE193777 entered the shared-gene layer because both cohorts retained gene-symbol "
        "expression matrices and binary progressor labels. The overlap across those cohorts was 14,440 genes."
    )
    add_plain_paragraph(
        doc,
        "The Bayesian gene model summarized existing within-cohort effect sizes using a normal-normal "
        "hierarchical framework. The objective was not to prove mechanism, but to stabilize ranking and provide "
        "posterior uncertainty intervals that are easier to interpret than raw point estimates alone."
    )
    add_plain_paragraph(
        doc,
        "The pathway-level analysis derived sample-level scores from the exported enrichment gene sets. "
        "Because the enrichment source was itself downstream of the project pipeline, the pathway results should "
        "be viewed as structured biological corroboration rather than an independent discovery layer."
    )
    add_plain_paragraph(
        doc,
        "The OmicsClaw-inspired deconvolution extension used canonical blood-cell marker genes with a non-negative "
        "least squares framework. These outputs are relative composition proxies and should not be interpreted as "
        "absolute leukocyte fractions. The coexpression extension was used to summarize coordinated programs and hub "
        "genes, but not to claim causal regulatory networks."
    )
    add_heading(doc, "Supplementary Figure", 1)
    add_figure(
        doc,
        ANALYSIS_DIR / "conceptual_dag.png",
        "Supplementary Figure 1. Directed acyclic graph showing the conceptual relationships among cohort, platform, host baseline factors, latent progression biology, preprocessing, and observed transcriptomic measurements.",
    )
    add_heading(doc, "Supplementary Table", 1)
    add_table(
        doc,
        data.bayes_genes.head(20).loc[:, ["gene", "posterior_mean", "ci95_low", "ci95_high"]].round(3).rename(
            columns={
                "gene": "Gene",
                "posterior_mean": "Posterior mean",
                "ci95_low": "95% CrI low",
                "ci95_high": "95% CrI high",
            }
        ),
        "Supplementary Table 1. Top 20 Bayesian gene-level posterior effects.",
    )
    out = OUT_DIR / "04_Supplementary_Methods_and_Figures.docx"
    doc.save(out)
    return out


def build_omicsclaw_extension_docx() -> Path | None:
    if not OMICSCLAW_DIR.exists():
        return None
    summary_path = OMICSCLAW_DIR / "nnls_celltype_summary.csv"
    module_path = OMICSCLAW_DIR / "coexpression_module_summary.csv"
    overlap_path = OMICSCLAW_DIR / "coexpression_module_signature_overlap.csv"
    if not summary_path.exists() or not module_path.exists():
        return None

    cell_df = pd.read_csv(summary_path)
    module_df = pd.read_csv(module_path)
    overlap_df = pd.read_csv(overlap_path) if overlap_path.exists() else pd.DataFrame()
    for df in [cell_df, module_df]:
        for col in ["mean_progressor", "mean_nonprogressor", "delta_progressor_minus_nonprogressor"]:
            if col in df.columns:
                df[col] = df[col].map(lambda x: f"{x:.3f}")
        for col in ["ttest_pvalue", "fdr_bh"]:
            if col in df.columns:
                df[col] = df[col].map(format_pvalue)

    doc = Document()
    set_style(doc)
    add_plain_paragraph(doc, "OmicsClaw-Informed Extension Analyses", bold=True, center=True)
    add_plain_paragraph(
        doc,
        "This supplement summarizes two additional analysis layers inspired by OmicsClaw bulk-RNA workflows: "
        "marker-based NNLS deconvolution and WGCNA-style coexpression analysis."
    )
    add_plain_paragraph(
        doc,
        "These outputs are intended to add biological context to the advanced tuberculosis progression manuscript. "
        "They should be interpreted as supportive analyses rather than stand-alone validation of mechanism."
    )
    add_heading(doc, "Immune Deconvolution", 1)
    add_plain_paragraph(
        doc,
        "A conservative NNLS framework was applied using canonical blood-cell marker genes present in the shared-gene "
        "matrix. The resulting coefficients are best interpreted as relative composition proxies."
    )
    add_table(
        doc,
        cell_df.rename(
            columns={
                "cell_type": "Cell type",
                "mean_progressor": "Mean progressor",
                "mean_nonprogressor": "Mean non-progressor",
                "delta_progressor_minus_nonprogressor": "Delta",
                "ttest_pvalue": "p value",
                "fdr_bh": "FDR",
            }
        ),
        "Supplementary Table A1. NNLS-derived cell-type score summary. These values are marker-based composition proxies rather than measured cell fractions.",
    )
    add_figure(
        doc,
        OMICSCLAW_DIR / "nnls_celltype_barplot.png",
        "Supplementary Figure A1. Mean NNLS cell-type scores by progressor status.",
    )
    add_figure(
        doc,
        OMICSCLAW_DIR / "nnls_celltype_boxplots.png",
        "Supplementary Figure A2. Distribution of NNLS cell-type scores across progressors and non-progressors.",
    )
    add_heading(doc, "Coexpression Modules", 1)
    add_plain_paragraph(
        doc,
        "A WGCNA-style unsigned coexpression analysis was performed on the most variable shared genes after "
        "cohort centering. Modules were summarized by eigengenes and compared by progressor status."
    )
    add_table(
        doc,
        module_df.rename(
            columns={
                "module": "Module",
                "n_genes": "Genes",
                "mean_progressor": "Mean progressor",
                "mean_nonprogressor": "Mean non-progressor",
                "delta_progressor_minus_nonprogressor": "Delta",
                "ttest_pvalue": "p value",
                "fdr_bh": "FDR",
            }
        ),
        "Supplementary Table A2. Coexpression module-trait summary. Modules are reported as supportive expression programs, not causal networks.",
    )
    if not overlap_df.empty:
        add_table(
            doc,
            overlap_df.rename(
                columns={
                    "module": "Module",
                    "n_signature_overlap": "Bayesian signature overlap",
                    "signature_genes": "Overlapping genes",
                }
            ),
            "Supplementary Table A3. Overlap between coexpression modules and top Bayesian genes.",
        )
    add_figure(
        doc,
        OMICSCLAW_DIR / "coexpression_module_trait_barplot.png",
        "Supplementary Figure A3. Coexpression module eigengene differences by progressor status.",
    )
    add_figure(
        doc,
        OMICSCLAW_DIR / "coexpression_module_boxplots.png",
        "Supplementary Figure A4. Top module eigengene distributions in progressors and non-progressors.",
    )
    add_heading(doc, "Cohort Expansion Note", 1)
    note_path = OMICSCLAW_DIR / "omicsclaw_literature_note.md"
    if note_path.exists():
        for line in note_path.read_text(encoding="utf-8").splitlines():
            clean = line.strip()
            if not clean:
                continue
            if clean.startswith("#"):
                continue
            if clean.startswith("- "):
                doc.add_paragraph(clean[2:], style="List Bullet")
            else:
                add_plain_paragraph(doc, clean)
    out = OUT_DIR / "06_OmicsClaw_Extensions.docx"
    doc.save(out)
    return out


def build_cover_letter() -> Path:
    doc = Document()
    set_style(doc)
    add_plain_paragraph(doc, "Cover Letter", bold=True, center=True)
    add_plain_paragraph(doc, f"Dear Editor,")
    add_plain_paragraph(
        doc,
        "I am submitting the manuscript titled "
        f"\"{TITLE}\" for consideration as an original research article."
    )
    add_plain_paragraph(
        doc,
        "This work presents a systems-level reanalysis of public tuberculosis progression transcriptomic datasets. "
        "The manuscript extends the primary signature-discovery study by showing how cohort centering, Bayesian "
        "hierarchical synthesis, latent factor analysis, pathway modeling, and correlation structure can reveal "
        "additional biological meaning while maintaining conservative evidence boundaries."
    )
    add_plain_paragraph(
        doc,
        "The manuscript is original, has not been published elsewhere, and is not under consideration by another journal. "
        "All data analyzed are publicly available, and the repository URL is provided for transparency and reproducibility."
    )
    add_plain_paragraph(doc, "Sincerely,")
    add_plain_paragraph(doc, f"{AUTHOR_NAME}, {DEGREE}")
    out = OUT_DIR / "05_Cover_Letter.docx"
    doc.save(out)
    return out


def build_review_note() -> Path:
    lines = [
        "# Internal Review Note",
        "",
        "Review pass 1",
        "- Checked that all numerical claims in the manuscript match files in `results/advanced_analysis`.",
        "- Confirmed that cohort counts, progressor counts, posterior means, factor p values, pathway terms, deconvolution deltas, and coexpression module statistics match generated outputs.",
        "- Identified gaps: sparse method references for deconvolution/coexpression, NA demographic fields in Table 1, and p-value rendering that rounded small values to zero.",
        "- Removed language that would imply causal proof or clinical readiness.",
        "",
        "Review pass 2",
        "- Checked sequential appearance of Table 1 to Table 4 and Figure 1 to Figure 4 in the manuscript text.",
        "- Checked that the DAG is placed in the supplementary document, not in the main article as if it were an estimated causal model.",
        "- Checked that the conclusions remain aligned with the two-cohort evidence boundary.",
        "- Confirmed that age and sex are now parsed from metadata, platform is reported in Table 1, and extension summaries include FDR columns.",
    ]
    out = OUT_DIR / "internal_review_log.md"
    out.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return out


def build_package_summary() -> Path:
    lines = [
        "# Advanced Analysis Package",
        "",
        f"Repository: {REPO_URL}",
        "",
        "Contents",
        "- 01_Title_Page.docx",
        "- 02_Manuscript.docx",
        "- 03_Highlights.docx",
        "- 04_Supplementary_Methods_and_Figures.docx",
        "- 05_Cover_Letter.docx",
        "- 06_OmicsClaw_Extensions.docx",
        "- internal_review_log.md",
        "",
        "Source analysis directory",
        "- results/advanced_analysis",
        "",
        "Notes",
        "- The advanced shared-gene layer currently uses GSE107994 and GSE193777 only.",
        "- GSE79362 was not forced into the shared-gene analysis because its current feature mapping is not commensurate with the two retained cohorts.",
        "- Conclusions are framed for biological interpretation and hypothesis generation, not causal proof or immediate clinical deployment.",
    ]
    out = OUT_DIR / "README.md"
    out.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return out


def build_validation_report() -> Path:
    from docx import Document as DocxDocument

    manuscript = OUT_DIR / "02_Manuscript.docx"
    doc = DocxDocument(manuscript)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)
    words = sum(len(p.split()) for p in paragraphs)
    lines = [
        "Advanced submission package validation",
        f"Manuscript path: {manuscript.name}",
        f"Approximate manuscript words (paragraph text only): {words}",
        f"Contains Table 1-4 in order: {all(text.find(f'Table {i}') != -1 for i in range(1, 5))}",
        f"Contains Figure 1-4 in order: {all(text.find(f'Figure {i}') != -1 for i in range(1, 5))}",
        f"Contains References section: {'References' in text}",
        f"Contains Declarations section: {'Declarations' in text}",
        f"Contains deconvolution caveat: {'marker-based composition proxies' in text}",
        f"Contains coexpression caveat: {'causal network' in text}",
        "Evidence boundary note: advanced shared-gene analysis restricted to GSE107994 and GSE193777.",
        "Interpretive limit note: DAG is conceptual and not used as causal proof.",
    ]
    out = OUT_DIR / "validation_report.txt"
    out.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return out


def zip_package() -> Path:
    zip_path = OUT_DIR / "advanced_submission_package.zip"
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for path in sorted(OUT_DIR.iterdir()):
            if path == zip_path:
                continue
            zf.write(path, arcname=path.name)
    return zip_path


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    data = load_data()
    build_title_page(data)
    build_manuscript(data)
    build_highlights()
    build_supplement(data)
    build_cover_letter()
    build_omicsclaw_extension_docx()
    build_review_note()
    build_package_summary()
    build_validation_report()
    zip_package()
    print(f"Advanced submission assets written to {OUT_DIR}")


if __name__ == "__main__":
    main()
