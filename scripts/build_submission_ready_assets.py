from __future__ import annotations

from datetime import datetime
from pathlib import Path
import textwrap
import zipfile

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "submission_ready" / "bmc_personalized_medicine_20260317"
REPO_URL = "https://github.com/hssling/tb-progression-transcriptome-meta"

AUTHOR = {
    "name": "Dr. Siddalingaiah H S, MD",
    "department": "Department of Community Medicine",
    "institution": "Shridevi Institute of Medical Sciences & Research Hospital",
    "city": "Tumkur",
    "state": "Karnataka",
    "country": "India",
    "email": "hssling@yahoo.com",
    "orcid": "0000-0002-4771-8285",
}

TITLE = (
    "A reproducible multi-cohort transcriptomic meta-analysis identifies "
    "host blood biomarkers associated with tuberculosis progression"
)
RUNNING_TITLE = "TB progression transcriptomic meta-analysis"


def set_default_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str = "", bold: bool = False, center: bool = False) -> None:
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold


def add_table_from_df(
    doc: Document,
    df: pd.DataFrame,
    caption: str,
    round_map: dict[str, int] | None = None,
    max_rows: int | None = None,
) -> None:
    add_para(doc, caption, bold=True)
    if max_rows is not None:
        df = df.head(max_rows)
    if round_map:
        df = df.copy()
        for col, digits in round_map.items():
            if col in df.columns:
                df[col] = df[col].map(lambda value: f"{value:.{digits}f}" if pd.notna(value) else "")
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for idx, col in enumerate(df.columns):
        table.cell(0, idx).text = str(col)
    for row in df.itertuples(index=False):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    doc.add_paragraph("")


def add_picture(doc: Document, path: Path, caption: str, width: float = 6.0) -> None:
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(width))
    p = doc.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, caption, center=True)


def load_inputs() -> dict[str, pd.DataFrame]:
    return {
        "dataset_summary": pd.read_csv(ROOT / "R_pipeline" / "output" / "dataset_summary.csv"),
        "meta_gene_list": pd.read_csv(ROOT / "submission_package" / "Tables" / "meta_gene_list.csv"),
        "validation_summary": pd.read_csv(ROOT / "R_pipeline" / "output" / "validation_summary.csv"),
        "loco_performance": pd.read_csv(ROOT / "results" / "tables" / "loco_performance.csv"),
        "pathway_enrichment": pd.read_csv(ROOT / "R_pipeline" / "output" / "pathway_enrichment_R.csv"),
        "lasso_signature": pd.read_csv(ROOT / "R_pipeline" / "output" / "lasso_signature.csv"),
        "pipeline_timings": pd.read_csv(ROOT / "submission_package" / "Tables" / "pipeline_timings.csv"),
    }


def build_title_page() -> Path:
    doc = Document()
    set_default_style(doc)
    add_para(doc, TITLE, bold=True, center=True)
    add_para(doc, "", center=True)
    add_para(doc, AUTHOR["name"], center=True)
    add_para(
        doc,
        (
            f'{AUTHOR["department"]}, {AUTHOR["institution"]}, '
            f'{AUTHOR["city"]}, {AUTHOR["state"]}, {AUTHOR["country"]}'
        ),
        center=True,
    )
    add_para(doc, f'ORCID: {AUTHOR["orcid"]}', center=True)
    add_para(doc, f'Correspondence: {AUTHOR["email"]}', center=True)
    add_para(doc, f"Running title: {RUNNING_TITLE}", center=True)
    add_para(doc, "Keywords: tuberculosis, incipient tuberculosis, transcriptomics, biomarker, meta-analysis", center=True)
    doc.add_page_break()
    add_heading(doc, "Declarations", 1)
    sections = {
        "Author details": (
            f'{AUTHOR["name"]}, {AUTHOR["department"]}, {AUTHOR["institution"]}, '
            f'{AUTHOR["city"]}, {AUTHOR["state"]}, {AUTHOR["country"]}.'
        ),
        "Funding": "No external funding was received for this study.",
        "Competing interests": "The author declares that there are no competing interests.",
        "Author contributions": (
            "Siddalingaiah H S conceived the study, curated the analysis framework, "
            "reviewed the computational outputs, drafted the manuscript, and approved the final version."
        ),
        "Availability of data and materials": (
            "Public-source cohort identifiers, generated summary tables, figures, manuscript assets, and "
            f"submission documentation are available in the project repository ({REPO_URL}). "
            "Primary transcriptomic studies were obtained from NCBI GEO."
        ),
        "Code availability": (
            "The reproducible codebase for discovery, harmonization, meta-analysis, reporting, and submission-asset "
            f"generation is available at {REPO_URL}."
        ),
        "Ethics approval and consent to participate": (
            "This work used de-identified publicly available transcriptomic datasets and did not involve new patient recruitment. "
            "Additional institutional ethics review was therefore not required for the present secondary analysis."
        ),
        "Consent for publication": "Not applicable.",
        "Acknowledgements": "The author acknowledges the original investigators who deposited the GEO datasets used in this analysis.",
    }
    for heading, text in sections.items():
        add_heading(doc, heading, 2)
        add_para(doc, text)
    out_path = OUT_DIR / "01_Title_Page_and_Declarations.docx"
    doc.save(out_path)
    return out_path


def build_main_article(data: dict[str, pd.DataFrame]) -> Path:
    doc = Document()
    set_default_style(doc)
    add_para(doc, TITLE, bold=True, center=True)
    add_para(doc, AUTHOR["name"], center=True)
    add_para(doc, "", center=True)

    add_heading(doc, "Abstract", 1)
    abstract = textwrap.dedent(
        """
        Background: Tuberculosis remains one of the leading infectious causes of death worldwide, and the ability to identify individuals who are progressing from latent or incipient infection toward active disease remains a major unmet need. Blood transcriptomic biomarkers are attractive because they are non-sputum based and biologically informative, but reproducibility across cohorts and assay platforms remains a central obstacle.

        Methods: We assembled a reproducible public-data workflow for tuberculosis progression transcriptomics using curated GEO studies. Three primary cohorts were processed in the current package, comprising 533 samples overall, while two cohorts contributed harmonized binary labels for strict leave-one-cohort-out validation. Gene-level effects were combined by random-effects meta-analysis, candidate genes were ranked by stability and meta z-score, and three classification approaches were assessed: elastic net, linear support vector machine, and a simple gene-set score. Functional interpretation used exported Gene Ontology enrichment outputs.

        Results: The analysis prioritized a 25-gene host signature associated with tuberculosis progression. Leading genes included MILR1, VSIG4, CCR2, CD36, FZD5, AQP1, CRISPLD2, and IRAK3. The strongest leave-one-cohort-out performance was observed for the gene-set-score model when GSE107994 was held out, with AUC-ROC 0.914 and AUC-PR 0.828. Mean AUC-ROC values were 0.884 for the gene-set-score model, 0.838 for linear SVM, and 0.780 for elastic net. Enrichment analysis emphasized angiogenesis-related remodeling, leukocyte-mediated immunity, myeloid activation, and phagocytosis.

        Conclusions: A transparent public-data workflow can recover a coherent host blood transcriptomic program associated with tuberculosis progression. However, the present evidence remains best interpreted as a reproducible biomarker-discovery resource rather than a locked clinical test because harmonized progression cohorts are limited and cross-platform external validation remains incomplete.
        """
    ).strip()
    for paragraph in abstract.split("\n\n"):
        add_para(doc, paragraph)

    add_heading(doc, "Introduction", 1)
    intro_paragraphs = [
        (
            "According to the World Health Organization Global Tuberculosis Report 2025, an estimated 10.7 million people developed tuberculosis in 2024 and 1.23 million died from the disease. These figures underline that earlier identification of people moving from latent infection toward active disease remains a major public health priority."
        ),
        (
            "The transition from latent or incipient infection to symptomatic tuberculosis is the most actionable stage for preventive intervention. A blood-based biomarker would be especially attractive because it could be deployed when sputum-based diagnostics are unavailable, insensitive, or impractical in asymptomatic high-risk contacts."
        ),
        (
            "Host transcriptomics has emerged as one of the most biologically compelling approaches to this problem. The ACS study by Zak et al. showed that a blood RNA signature could anticipate future tuberculosis, while Singhania et al. demonstrated that tuberculosis infection states are transcriptionally heterogeneous rather than cleanly binary. More recent work in household contacts from Southern India further strengthened the case that host blood biosignatures can capture subclinical and incipient disease biology."
        ),
        (
            "Despite that progress, reproducibility remains the field's central weakness. Published signatures often lose performance when transferred across cohorts, geographies, platforms, or preprocessing pipelines. Some of that attenuation reflects genuine biological heterogeneity, but some reflects preventable methodological issues such as inconsistent gene mapping, weak harmonization, and data leakage caused by random row-wise validation strategies that ignore study boundaries."
        ),
        (
            "Public datasets are therefore both an opportunity and a stress test. They provide independent cohorts collected under different conditions, but they also expose how difficult it is to make a signature travel across studies without inflating claims. Public-data biomarker work is only useful if it is explicit about what was actually harmonized, what was only identified as a candidate resource, and what level of validation the resulting model genuinely achieved."
        ),
        (
            "The aim of the present study was to construct a reproducible workflow spanning cohort curation, harmonization, gene-level meta-analysis, machine learning validation, pathway interpretation, and manuscript generation. We sought to identify a coherent host blood signature associated with tuberculosis progression while keeping the final claims aligned with the actual support provided by the processed cohorts and exported outputs."
        ),
    ]
    for paragraph in intro_paragraphs:
        add_para(doc, paragraph)

    add_heading(doc, "Methods", 1)
    method_sections = {
        "Study design and data sources": (
            "This study used a reproducible public-data workflow centered on GEO studies relevant to host blood transcriptomic signatures of tuberculosis progression, incipient tuberculosis, subclinical tuberculosis, or closely related progression-focused designs. Eligible studies were restricted to human whole-blood or PBMC transcriptomic datasets with usable phenotype annotations and sufficient expression data to support downstream harmonization."
        ),
        "Cohort curation": (
            "The curated registry identified six relevant public studies. Three cohorts were retained as primary processed cohorts in the current package: GSE107994, representing the Leicester cohort; GSE193777, representing adult household contacts from Southern India; and GSE79362, corresponding to the ACS prospective cohort. Additional microarray studies, including GSE19491, GSE37250, and GSE39940, were catalogued as candidate cohorts for future expansion but were not advanced through the same finalized validation path in this release."
        ),
        "Preprocessing and harmonization": (
            "Expression matrices were normalized and represented at gene level wherever possible. Duplicate or ambiguous mappings were collapsed conservatively to avoid inflated feature counts, after which datasets were restricted to a shared feature space prior to meta-analysis and model assessment. The current manuscript intentionally describes only the harmonized space supported by the exported results rather than implying universal compatibility across every identified cohort."
        ),
        "Meta-analysis": (
            "Per-gene effects were combined using a random-effects framework, and candidate progression-associated genes were ranked by absolute meta z-score. False-discovery adjustment and heterogeneity metrics were retained in the output tables. Because only two cohorts currently contributed directly comparable binary labels to the main validation framework, the meta-analysis was interpreted as a prioritization strategy for robust candidate genes rather than as proof of a universally stable clinical signature."
        ),
        "Model development and validation": (
            "Model assessment was built around leave-one-cohort-out validation to reduce information leakage across studies. Three approaches were evaluated: elastic net logistic regression, linear support vector machine, and a simple gene-set score. Performance was summarized using AUC-ROC, AUC-PR, and Brier score. This validation strategy was deliberately harsher than row-wise random splitting because the purpose of the analysis was to test whether a signal could survive true cohort transfer."
        ),
        "Pathway analysis": (
            "The leading genes were interpreted with the exported Gene Ontology biological-process enrichment tables generated by the R reporting pipeline. Only terms directly present in the exported results were used in the narrative. This restriction was intentional and was meant to prevent retrospective over-interpretation of the biology."
        ),
        "Reporting strategy": (
            "The present article was generated as part of a submission-ready reporting workflow that also produces the title page, declarations, supplementary information, cover letter, and journal-targeting notes from version-controlled outputs. This reporting layer was designed to reduce the gap between computational analysis and publication preparation while preserving a clear audit trail back to the repository and exported results."
        ),
    }
    for heading, text in method_sections.items():
        add_heading(doc, heading, 2)
        add_para(doc, text)

    add_heading(doc, "Results", 1)
    add_heading(doc, "Curated cohort set", 2)
    add_para(
        doc,
        "The curated registry identified six relevant public studies, of which three were processed as primary cohorts in the current run. These three studies span distinct epidemiological and technical settings, including a Leicester cohort, a Southern India household-contact cohort, and the ACS prospective cohort."
    )
    add_para(
        doc,
        "Across the primary processed cohorts, 533 samples were represented in the exported summaries. The cohort structures were not interchangeable: some were directly suitable for binary progressor versus nonprogressor modeling, whereas others primarily served as discovery or contextual resources. This asymmetry is a practical constraint in public-data tuberculosis biomarker work and explains why apparently large public resources may still yield only a limited set of fully harmonized validation-ready comparisons."
    )
    cohort_table = data["dataset_summary"].loc[
        data["dataset_summary"]["Status"].str.contains("Primary", na=False),
        ["GEO_ID", "Platform", "Total_Samples", "Groups", "PMID"],
    ]
    add_table_from_df(doc, cohort_table, "Table 1. Primary cohorts carried into the current submission package.")

    add_heading(doc, "Progression-associated genes", 2)
    add_para(
        doc,
        "The random-effects meta-analysis prioritized genes linked to innate immune regulation, phagocytic biology, leukocyte trafficking, and vascular remodeling. MILR1 was the top-ranked signal, followed closely by VSIG4, CCR2, CD36, and FZD5. Several of these genes have plausible mechanistic links to myeloid activation, immune-cell recruitment, tissue remodeling, and counter-regulatory signaling, making the overall ranking biologically coherent rather than a disconnected list of nominally significant transcripts."
    )
    add_para(
        doc,
        "The exported effect table showed strong positive meta effects for MILR1, VSIG4, CCR2, CD36, FZD5, and AQP1, while genes such as EPN2 and PLD4 showed inverse directionality within the leading ranks. IRAK3 remained among the most strongly associated positive genes and is notable because it encodes a negative regulator of Toll-like receptor signaling. In the setting of tuberculosis progression, this pattern is compatible with a host-response state that combines inflammatory recruitment with regulatory dampening rather than a simple monotonic activation program."
    )
    add_para(
        doc,
        "The effect sizes in the leading set were not only statistically compelling but also directionally coherent across the limited number of directly comparable cohorts in the current binary analysis. That point matters because robustness across cohorts is more relevant than extremity within a single study when the intended downstream use is cross-setting risk stratification."
    )
    gene_table = data["meta_gene_list"].loc[:, ["gene", "meta_effect", "meta_z", "meta_fdr", "i2"]]
    add_table_from_df(
        doc,
        gene_table,
        "Table 2. Top 15 progression-associated genes from the random-effects meta-analysis.",
        round_map={"meta_effect": 3, "meta_z": 3, "meta_fdr": 3, "i2": 1},
        max_rows=15,
    )

    add_heading(doc, "Cross-cohort validation", 2)
    add_para(
        doc,
        "Cross-cohort validation was strongest for the gene-set-score model, which produced the best held-out discrimination in GSE107994, with AUC-ROC 0.914 and AUC-PR 0.828. Linear SVM produced somewhat lower ranking performance but a more favorable Brier score profile, suggesting better calibration behavior in the current setting. Elastic net showed the greatest instability across the two held-out cohorts."
    )
    add_para(
        doc,
        "The validation table illustrates an important translational point: the simplest model was not necessarily the weakest. The gene-set-score approach retained the strongest cross-cohort ranking discrimination even though it did not yield the best calibration. That suggests that the underlying signal is real, but the transformation of that signal into calibrated individual risk estimates will require broader cohort coverage and better prevalence anchoring than is currently available."
    )
    add_para(
        doc,
        "Performance differences between GSE107994 and GSE193777 also highlight the role of cohort context. The same signature behaved more strongly in one held-out dataset than in the other, which is the kind of result that should be expected in public multi-study biomarker work. Rather than weakening the manuscript, this pattern strengthens the argument for honest external validation and against overfitting-prone discovery strategies."
    )
    perf_table = data["loco_performance"].copy()
    add_table_from_df(
        doc,
        perf_table,
        "Table 3. Leave-one-cohort-out model performance.",
        round_map={"auc_roc": 3, "auc_pr": 3, "brier": 3},
    )
    add_para(
        doc,
        "The exported validation summary further reported mean AUC-ROC values of 0.884 for the gene-set-score model, 0.838 for linear SVM, and 0.780 for elastic net, with the best left-out cohort being GSE107994. These values support the claim that a progression-related host signal is recoverable across independent cohorts, even though the present package does not yet support a definitive locked model for prospective clinical use."
    )

    add_heading(doc, "Functional interpretation", 2)
    add_para(
        doc,
        "Enrichment analysis highlighted regulation of angiogenesis, vasculature development, immune effector processes, leukocyte degranulation, myeloid activation, and phagocytosis. These findings imply that the leading progression-associated genes do not simply reflect a generic inflammatory state; rather, they point to remodeling of the host immune microenvironment, including cell trafficking, endothelial interaction, effector regulation, and phagocytic handling."
    )
    add_para(
        doc,
        "Several of the leading enriched terms were driven by recurring genes such as CCR2, CD36, FZD5, ITGB2, and SPARC. This recurrence is useful because it ties together the individual gene ranking and pathway-level interpretation. CCR2 is relevant to chemotactic recruitment, CD36 is linked to lipid handling and immune recognition, ITGB2 marks leukocyte adhesion and trafficking, and SPARC is implicated in tissue remodeling."
    )
    add_para(
        doc,
        "The prominence of IRAK3 adds a further mechanistic layer because it is a negative regulator of Toll-like receptor signaling and may mark a state of dysregulated feedback during incipient disease. Likewise, MILR1 suggests that mast-cell-related inhibitory signaling may deserve more attention than it has received in mainstream tuberculosis transcriptomic discussions."
    )
    pathway_table = data["pathway_enrichment"].loc[:, ["ID", "Description", "GeneRatio", "p.adjust", "geneID"]]
    add_table_from_df(
        doc,
        pathway_table,
        "Table 4. Representative enriched biological processes among leading progression-associated genes.",
        round_map={"p.adjust": 4},
        max_rows=12,
    )

    add_heading(doc, "Discussion", 1)
    discussion_paragraphs = [
        (
            "The central finding of this study is that a coherent host blood transcriptomic program associated with tuberculosis progression can be recovered from public datasets when the workflow is explicit, cohort boundaries are respected during validation, and the narrative is limited to what the processed data actually support. In that sense, the value of the present work lies not only in the gene list itself but also in the reproducible path by which the list was generated, summarized, and prepared for submission."
        ),
        (
            "The leading genes form a plausible biological core. MILR1 emerged as the most prominent signal and is notable because it points beyond the most frequently discussed interferon-centered views of host response. VSIG4, CCR2, CD36, FZD5, AQP1, and IRAK3 reinforce a picture of altered myeloid activation, cell adhesion, immune trafficking, and regulatory feedback. This pattern is compatible with an incipient-disease state in which host defense, immune modulation, and tissue remodeling are occurring simultaneously."
        ),
        (
            "The present findings align with the broader progression-signature literature while remaining distinct from it. The ACS work by Zak et al. established the principle that blood RNA can anticipate disease risk. Singhania et al. showed that host transcriptional states are heterogeneous rather than binary. Rajamanickam et al. extended this perspective to subclinical and incipient disease in household contacts. The current workflow adds a complementary contribution by emphasizing reproducible cross-cohort evaluation and integrating the reporting and submission pipeline into the analytical workflow itself."
        ),
        (
            "A major strength of the present approach is the refusal to treat random within-study validation as sufficient evidence. In multi-cohort transcriptomics, row-wise random splitting can allow study-specific artifacts to leak into both training and testing sets, inflating performance estimates. By contrast, leave-one-cohort-out validation is harsher but more informative. The fact that the signature retained meaningful discriminatory ability under that framework suggests that the observed signal is not purely cohort-specific."
        ),
        (
            "The limitations remain substantial and must stay central to interpretation. Only two cohorts presently support the harmonized binary validation framework used for the main performance claims. Additional candidate microarray studies have been identified but are not yet integrated into the same finalized validation path. The Brier scores indicate that ranking performance is better than calibration, meaning the current models are better understood as discriminative tools than as directly deployable risk calculators. Public metadata also do not consistently preserve the time-to-event resolution needed for stronger clinical-window claims."
        ),
        (
            "These limitations shape the translational meaning of the manuscript. The current package supports a methods-forward biomarker-discovery claim and offers a useful ranked candidate set for further investigation, including qRT-PCR panel development and prospective cohort validation. It does not yet justify a strong claim that the 25-gene signature is ready for immediate clinical implementation. That distinction is not a weakness; it is the correct interpretation of the available data."
        ),
        (
            "The pathway results help explain why translation may be both promising and difficult. The dominant themes are not narrowly limited to one canonical interferon signature. Instead, the enriched processes indicate immune effector regulation, leukocyte adhesion and degranulation, myeloid activation, phagocytosis, and vascular remodeling. Such a host-response mixture may generalize better than a single-axis inflammatory program, but it may also vary by clinical context, comorbidity, and specimen processing."
        ),
        (
            "Future work should proceed in three directions. First, additional progression cohorts, particularly those with prospective time-to-disease labeling, should be integrated into the same harmonization framework. Second, calibration-focused modeling should be revisited after broader cohort coverage, potentially combining transcriptomic signals with clinical covariates. Third, assay translation should prioritize genes that are not only statistically strong but also biologically interpretable, technically stable, and feasible for multiplex measurement."
        ),
        (
            "From a publication perspective, this manuscript is best positioned as a reproducible translational bioinformatics study rather than a definitive diagnostic benchmark. That framing is appropriate for the evidence generated here and avoids the overstatement that has contributed to skepticism around some previous host-signature reports. The package is intended to narrow the gap between exploratory public-data analysis and journal-ready reporting without disguising the uncertainty that still remains."
        ),
    ]
    for paragraph in discussion_paragraphs:
        add_para(doc, paragraph)

    add_heading(doc, "Conclusions", 1)
    add_para(
        doc,
        "A reproducible public-data workflow identified a biologically coherent host blood transcriptomic signature associated with tuberculosis progression risk and converted that workflow into a submission-ready manuscript package. The leading genes, including MILR1, VSIG4, CCR2, CD36, FZD5, AQP1, and IRAK3, support a model of progression characterized by coordinated innate immune regulation, leukocyte trafficking, phagocytic remodeling, and regulatory feedback."
    )
    add_para(
        doc,
        "The cross-cohort validation results indicate that this signal is not confined to a single study, but the current evidence remains best understood as a robust discovery resource rather than a finalized clinical test. Additional harmonized cohorts, better-calibrated models, and prospective external validation are the necessary next steps for translational advancement."
    )

    add_heading(doc, "Figures", 1)
    figure_dir = ROOT / "R_pipeline" / "figures"
    add_picture(
        doc,
        figure_dir / "forest_plot.png",
        "Figure 1. Forest-style summary plot of progression-associated transcriptomic effects across included cohorts.",
    )
    add_picture(
        doc,
        figure_dir / "roc_combined.png",
        "Figure 2. Combined ROC comparison across the evaluated validation models.",
    )
    add_picture(
        doc,
        figure_dir / "heatmap_signature.png",
        "Figure 3. Heatmap of the leading gene signature across the harmonized cohort set.",
    )
    add_picture(
        doc,
        figure_dir / "go_bp_dotplot.png",
        "Figure 4. Gene Ontology biological-process enrichment profile for the leading signature genes.",
    )

    add_heading(doc, "Declarations", 1)
    declaration_items = [
        ("Ethics approval and consent to participate", "No new human participants were enrolled. The analysis used de-identified public data only."),
        ("Consent for publication", "Not applicable."),
        ("Availability of data and materials", f"All generated submission assets, manuscript sources, figures, and tabular outputs are documented in the Git repository: {REPO_URL}."),
        ("Code availability", f"The full codebase used to prepare this analysis is maintained at {REPO_URL}."),
        ("Competing interests", "The author declares no competing interests."),
        ("Funding", "No external funding was received."),
        ("Author contributions", "Siddalingaiah H S performed the study design, analysis oversight, manuscript preparation, and final approval."),
    ]
    for heading, text in declaration_items:
        add_heading(doc, heading, 2)
        add_para(doc, text)

    add_heading(doc, "References", 1)
    references = [
        "World Health Organization. Global Tuberculosis Report 2025. Geneva: World Health Organization; 2025.",
        "Houben RMGJ, Dodd PJ. The global burden of latent tuberculosis infection: a re-estimation using mathematical modelling. PLoS Medicine. 2016;13:e1002152.",
        "Zak DE, Penn-Nicholson A, Scriba TJ, et al. A blood RNA signature for tuberculosis disease risk: a prospective cohort study. Lancet. 2016;387:2312-2322.",
        "Singhania A, Verma R, Graham CM, et al. A modular transcriptional signature identifies phenotypic heterogeneity of human tuberculosis infection. Nature Communications. 2018;9:2308.",
        "Rajamanickam A, Munisankar S, Dolla CK, et al. Host blood-based biosignatures for subclinical TB and incipient TB: a prospective study of adult TB household contacts in Southern India. Frontiers in Immunology. 2022;13.",
        "Sweeney TE, Braviak L, Tato CM, Khatri P. Genome-wide expression for diagnosis of pulmonary tuberculosis: a multicohort analysis. Lancet Respiratory Medicine. 2016;4:213-224.",
        "Warsinske HC, Rao AM, Moreira FMF, et al. Assessment of validity of a blood-based 3-gene signature score for progression and diagnosis of tuberculosis, disease severity, and treatment response. JAMA Network Open. 2018;1:e183779.",
        "Davis S, Meltzer PS. GEOquery: a bridge between the Gene Expression Omnibus and BioConductor. Bioinformatics. 2007;23:1846-1847.",
        "Ritchie ME, Phipson B, Wu D, et al. limma powers differential expression analyses for RNA-sequencing and microarray studies. Nucleic Acids Research. 2015;43:e47.",
        "Leek JT, Johnson WE, Parker HS, Jaffe AE, Storey JD. The sva package for removing batch effects and other unwanted variation in high-throughput experiments. Bioinformatics. 2012;28:882-883.",
        "Balduzzi S, Rucker G, Schwarzer G. How to perform a meta-analysis with R: a practical tutorial. Evidence-Based Mental Health. 2019;22:153-160.",
        "Friedman J, Hastie T, Tibshirani R. Regularization paths for generalized linear models via coordinate descent. Journal of Statistical Software. 2010;33:1-22.",
        "Wu T, Hu E, Xu S, et al. clusterProfiler 4.0: a universal enrichment tool for interpreting omics data. Innovation. 2021;2:100141.",
        "Wesche H, Gao X, Li X, et al. IRAK-M is a novel member of the interleukin-1 receptor-associated kinase family. Journal of Biological Chemistry. 1999;274:19403-19410.",
        "The tb-progression-transcriptome-meta repository. Available at: https://github.com/hssling/tb-progression-transcriptome-meta.",
    ]
    for ref in references:
        add_para(doc, ref)

    out_path = OUT_DIR / "02_Main_Article.docx"
    doc.save(out_path)
    return out_path


def build_supplement(data: dict[str, pd.DataFrame]) -> Path:
    doc = Document()
    set_default_style(doc)
    add_para(doc, "Supplementary Information", bold=True, center=True)
    add_para(doc, TITLE, center=True)

    add_heading(doc, "Supplementary Methods", 1)
    supplement_text = [
        "The computational workflow combined Python-based project orchestration with an R-based downstream reporting layer. The package preserved exported cohort summaries, meta-analysis tables, LOCO validation metrics, pathway enrichment results, and manuscript assets in version-controlled form.",
        "Primary processed cohorts were the Leicester cohort (GSE107994), the Southern India household-contact cohort (GSE193777), and the ACS cohort (GSE79362). Candidate microarray cohorts remained catalogued for future harmonization.",
        "The current supplementary file emphasises reproducibility inventory rather than restating every package dependency already captured in the repository.",
    ]
    for paragraph in supplement_text:
        add_para(doc, paragraph)

    add_heading(doc, "Supplementary Table S1. Stability-ranked minimal signature", 1)
    add_table_from_df(
        doc,
        data["lasso_signature"].loc[:, ["rank", "gene", "stability", "meta_z"]],
        "Stability-guided ranking of minimal signature candidates.",
        round_map={"stability": 2, "meta_z": 3},
        max_rows=15,
    )

    add_heading(doc, "Supplementary Table S2. Validation summary", 1)
    add_table_from_df(doc, data["validation_summary"], "Exported validation summary from the reporting pipeline.")

    add_heading(doc, "Supplementary Table S3. Pipeline timings", 1)
    add_table_from_df(doc, data["pipeline_timings"], "End-to-end pipeline timing summary.", round_map={"elapsed_seconds": 1})

    add_heading(doc, "Supplementary Figure Notes", 1)
    notes = [
        "Figure 1 source: R_pipeline/figures/forest_plot.png",
        "Figure 2 source: R_pipeline/figures/roc_combined.png",
        "Figure 3 source: R_pipeline/figures/heatmap_signature.png",
        "Figure 4 source: R_pipeline/figures/go_bp_dotplot.png",
        "Supplementary assets are reproducibly linked to the repository and the output package in this folder.",
    ]
    for note in notes:
        add_para(doc, note)

    out_path = OUT_DIR / "03_Supplementary_Information.docx"
    doc.save(out_path)
    return out_path


def build_cover_letter() -> Path:
    doc = Document()
    set_default_style(doc)
    add_para(doc, datetime.now().strftime("%d %B %Y"))
    add_para(doc, "The Editor")
    add_para(doc, "BMC Personalized Medicine")
    add_para(doc, "Springer Nature")
    add_para(doc, "")
    add_para(doc, "Re: Submission of an original research article")
    add_para(doc, "")
    body = [
        f'I am submitting the manuscript titled "{TITLE}" for consideration as an Original Research article in BMC Personalized Medicine.',
        "This study presents a reproducible public-data workflow for identifying host blood transcriptomic biomarkers associated with tuberculosis progression. "
        "The package combines cohort curation, gene-level meta-analysis, leave-one-cohort-out validation, pathway interpretation, and a submission-ready documentation layer.",
        "The manuscript is suited to the journal because it focuses on biomarker prioritization, cross-cohort reproducibility, and translational interpretation rather than a narrow single-cohort discovery claim. "
        "The submission is especially appropriate for a methods-forward personalized-risk framing in infectious disease.",
        "This manuscript has not been published elsewhere, is not under consideration by another journal, and uses only de-identified publicly available datasets. "
        "All code, figures, tables, and documentation are available in the associated GitHub repository.",
        f"Repository: {REPO_URL}",
        f"Corresponding author: {AUTHOR['name']} ({AUTHOR['email']})",
    ]
    for paragraph in body:
        add_para(doc, paragraph)
    add_para(doc, "")
    add_para(doc, "Sincerely,")
    add_para(doc, AUTHOR["name"])
    add_para(doc, AUTHOR["department"])
    add_para(doc, AUTHOR["institution"])
    add_para(doc, f'{AUTHOR["city"]}, {AUTHOR["state"]}, {AUTHOR["country"]}')
    out_path = OUT_DIR / "04_Cover_Letter.docx"
    doc.save(out_path)
    return out_path


def build_journal_notes() -> Path:
    doc = Document()
    set_default_style(doc)
    add_para(doc, "Journal Fit and Submission Notes", bold=True, center=True)

    add_heading(doc, "Primary target journal", 1)
    add_para(
        doc,
        "BMC Personalized Medicine (Springer Nature / BMC) is the primary recommended target for this package because "
        "the manuscript is positioned as a reproducible biomarker-discovery and risk-stratification resource rather than a definitive clinical assay paper."
    )
    add_para(
        doc,
        "This package was therefore formatted as a BMC-style original research submission with a separate title page, main article, supplementary file, and cover letter."
    )

    add_heading(doc, "Cost position", 1)
    add_para(
        doc,
        "As checked on 17 March 2026 from official Springer Nature / BMC material, BMC Personalized Medicine is operating under a launch-period arrangement in which publication costs are covered until 31 December 2026. "
        "This makes it the most suitable APC-free option within the Springer Nature family for the present manuscript."
    )

    add_heading(doc, "Fallback subscription-format option", 1)
    add_para(
        doc,
        "If a more conventional subscription or hybrid journal is preferred, Infection (Springer) is the recommended fallback. "
        "It is a clinically oriented infectious-disease journal within Springer and uses Springer hybrid publishing options, so a non-open-access route remains possible."
    )

    add_heading(doc, "Repository and audit trail", 1)
    add_para(doc, f"GitHub repository: {REPO_URL}")
    add_para(
        doc,
        "All submission files in this folder were generated from repository outputs and should be cited together with the repository when the manuscript is uploaded or shared."
    )

    out_path = OUT_DIR / "05_Journal_Fit_and_Submission_Notes.docx"
    doc.save(out_path)
    return out_path


def build_readme() -> Path:
    text = textwrap.dedent(
        f"""
        # Submission-ready package

        Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}

        Primary target journal: BMC Personalized Medicine (Springer Nature)
        Fallback journal: Infection (Springer)
        Repository: {REPO_URL}

        Files:
        - 01_Title_Page_and_Declarations.docx
        - 02_Main_Article.docx
        - 03_Supplementary_Information.docx
        - 04_Cover_Letter.docx
        - 05_Journal_Fit_and_Submission_Notes.docx
        - submission_ready_package.zip

        Notes:
        - The manuscript is intentionally framed as a reproducible biomarker-discovery resource.
        - Figures are embedded directly in the main article DOCX.
        - Supplementary tables are included in the supplementary DOCX.
        - The package cites the GitHub repository and uses the current local author metadata.
        """
    ).strip() + "\n"
    out_path = OUT_DIR / "README.md"
    out_path.write_text(text, encoding="utf-8")
    return out_path


def zip_outputs(paths: list[Path]) -> Path:
    zip_path = OUT_DIR / "submission_ready_package.zip"
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for path in paths:
            zf.write(path, arcname=path.name)
    return zip_path


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    data = load_inputs()
    outputs = [
        build_title_page(),
        build_main_article(data),
        build_supplement(data),
        build_cover_letter(),
        build_journal_notes(),
        build_readme(),
    ]
    outputs.append(zip_outputs(outputs))
    print("Generated submission-ready assets:")
    for path in outputs:
        print(path.relative_to(ROOT))


if __name__ == "__main__":
    main()
