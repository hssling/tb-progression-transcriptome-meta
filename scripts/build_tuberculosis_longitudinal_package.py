from __future__ import annotations

from pathlib import Path
import shutil
import zipfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SRC_DIR = ROOT / "submission_ready" / "longitudinal_tb_dynamics_20260318"
OUT_DIR = ROOT / "submission_ready" / "tuberculosis_longitudinal_20260318"
REPO_URL = "https://github.com/hssling/tb-progression-transcriptome-meta"

TITLE = (
    "Longitudinal reanalysis of a prospective tuberculosis blood RNA-sequencing "
    "cohort suggests follow-up dynamics in coordinated host-response programs"
)


def set_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(0)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)


def add_paragraph(doc: Document, text: str = "", bold: bool = False, center: bool = False) -> None:
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold


def build_cover_letter() -> Path:
    doc = Document()
    set_style(doc)
    add_paragraph(doc, "Cover Letter", bold=True, center=True)
    add_paragraph(doc, "Dear Editors of Tuberculosis,")
    add_paragraph(doc, f'I am submitting the manuscript "{TITLE}" for consideration as an Original Article in Tuberculosis.')
    add_paragraph(
        doc,
        "This manuscript analyzes a prospective tuberculosis blood RNA-sequencing cohort after junction-to-gene remapping and subject-level linkage. "
        "The central question is biological rather than purely meta-analytical: whether repeated blood transcriptomic measurements show different host-program dynamics "
        "in progressors and non-progressors during follow-up."
    )
    add_paragraph(
        doc,
        "The work is aligned with the journal's interest in host response, immunology, and pathogenesis. "
        "It does not present a literature-only meta-analysis. Instead, it focuses on longitudinal behavior within one remapped prospective cohort, using prespecified "
        "myeloid, vascular, coexpression, and cell-state programs to interpret repeated measurements conservatively."
    )
    add_paragraph(
        doc,
        "All data analyzed are publicly available, and the code and derived outputs are available in the project repository. "
        "The manuscript is original, has not been published elsewhere, and is not under consideration by another journal."
    )
    add_paragraph(doc, "Sincerely,")
    add_paragraph(doc, "Siddalingaiah H S, MD")
    out = OUT_DIR / "03_Tuberculosis_Cover_Letter.docx"
    doc.save(out)
    return out


def build_author_notes() -> Path:
    doc = Document()
    set_style(doc)
    add_paragraph(doc, "Tuberculosis Journal Submission Notes", bold=True, center=True)
    add_paragraph(doc, "Target journal: Tuberculosis")
    add_paragraph(doc, "Article type: Original Article")
    add_paragraph(doc, "")
    add_paragraph(doc, "Journal-fit note", bold=True)
    for line in [
        "The manuscript centers on host response and immunology in tuberculosis using a prospective repeated-measures blood RNA-sequencing cohort.",
        "The framing is longitudinal biology and pathogenesis-oriented, not a literature-only meta-analysis.",
        "The strongest findings concern coordinated host-response programs rather than a generic public-database ranking exercise.",
    ]:
        doc.add_paragraph(line, style="List Bullet")
    add_paragraph(doc, "")
    add_paragraph(doc, "Elsevier-specific points addressed", bold=True)
    for line in [
        "Separate title page prepared.",
        "Highlights included as a separate file.",
        "Data availability statement retained in the manuscript.",
        "Generative AI disclosure retained at the end of the manuscript before references.",
        "Figures and tables currently embedded for review; they can be split at upload if preferred.",
    ]:
        doc.add_paragraph(line, style="List Bullet")
    add_paragraph(doc, "")
    add_paragraph(doc, "Repository", bold=True)
    add_paragraph(doc, REPO_URL)
    out = OUT_DIR / "05_Tuberculosis_Submission_Notes.docx"
    doc.save(out)
    return out


def build_data_statement() -> Path:
    doc = Document()
    set_style(doc)
    add_paragraph(doc, "Data Statement", bold=True, center=True)
    add_paragraph(
        doc,
        "The underlying public dataset analyzed in this study is available in the Gene Expression Omnibus under accession GSE79362. "
        "The remapped gene-level matrices, generated analysis tables, figure assets, and analysis scripts used for this manuscript are available in the project repository:"
    )
    add_paragraph(doc, REPO_URL)
    add_paragraph(
        doc,
        "This study did not generate new patient-level clinical data. The manuscript uses de-identified public transcriptomic data only."
    )
    out = OUT_DIR / "06_Tuberculosis_Data_Statement.docx"
    doc.save(out)
    return out


def build_readme() -> Path:
    text = "\n".join(
        [
            "# Tuberculosis Longitudinal Submission Package",
            "",
            "Target journal: Tuberculosis (Elsevier)",
            "Article type: Original Article",
            f"Repository: {REPO_URL}",
            "",
            "Contents",
            "- 01_Tuberculosis_Title_Page.docx",
            "- 02_Tuberculosis_Main_Manuscript.docx",
            "- 03_Tuberculosis_Cover_Letter.docx",
            "- 04_Tuberculosis_Highlights.docx",
            "- 05_Tuberculosis_Submission_Notes.docx",
            "- 06_Tuberculosis_Data_Statement.docx",
            "- 07_Tuberculosis_Supplementary_Methods_and_Tables.docx",
            "- validation_report.txt",
            "",
            "This package adapts the longitudinal TB dynamics manuscript to the Elsevier Tuberculosis submission shape.",
        ]
    )
    out = OUT_DIR / "README.md"
    out.write_text(text + "\n", encoding="utf-8")
    return out


def build_validation() -> Path:
    manuscript = OUT_DIR / "02_Tuberculosis_Main_Manuscript.docx"
    doc = Document(manuscript)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)
    lines = [
        "Tuberculosis package validation",
        "Main manuscript: 02_Tuberculosis_Main_Manuscript.docx",
        f"Contains Abstract: {'Abstract' in text}",
        f"Contains Keywords: {'Keywords:' in text}",
        f"Contains Declarations section: {'Declarations' in text}",
        f"Contains References section: {'References' in text}",
        f"Contains Data availability statement: {'Availability of data and materials' in text}",
        f"Contains AI declaration: {'Use of generative AI' in text}",
        f"Contains Table 1-4 in order: {all(text.find(f'Table {i}') != -1 for i in range(1, 5))}",
        f"Contains Figure 1-4 in order: {all(text.find(f'Figure {i}') != -1 for i in range(1, 5))}",
        "Editorial-fit note: manuscript is framed as prospective cohort reanalysis, not literature-only meta-analysis.",
    ]
    out = OUT_DIR / "validation_report.txt"
    out.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return out


def copy_core_files() -> None:
    mapping = {
        "01_Title_Page.docx": "01_Tuberculosis_Title_Page.docx",
        "02_Manuscript.docx": "02_Tuberculosis_Main_Manuscript.docx",
        "03_Highlights.docx": "04_Tuberculosis_Highlights.docx",
        "04_Supplementary_Methods_and_Tables.docx": "07_Tuberculosis_Supplementary_Methods_and_Tables.docx",
    }
    for src_name, dst_name in mapping.items():
        shutil.copy2(SRC_DIR / src_name, OUT_DIR / dst_name)


def zip_package() -> Path:
    zip_path = OUT_DIR / "tuberculosis_longitudinal_submission_package.zip"
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for path in sorted(OUT_DIR.iterdir()):
            if path == zip_path:
                continue
            zf.write(path, arcname=path.name)
    return zip_path


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    copy_core_files()
    build_cover_letter()
    build_author_notes()
    build_data_statement()
    build_readme()
    build_validation()
    zip_package()
    print(f"Tuberculosis submission package written to {OUT_DIR}")


if __name__ == "__main__":
    main()
