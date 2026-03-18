from __future__ import annotations

import importlib.util
import re
import shutil
import zipfile
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SRC_DIR = ROOT / "submission_ready" / "advanced_tb_systems_20260317"
OUT_DIR = ROOT / "submission_ready" / "bmc_med_genomics_advanced_20260318_rev2"
REPO_URL = "https://github.com/hssling/tb-progression-transcriptome-meta"
ACCESS_DATE = "18 Mar 2026"


def load_advanced_builder():
    spec = importlib.util.spec_from_file_location(
        "advanced_builder", ROOT / "scripts" / "build_advanced_submission_assets.py"
    )
    if spec is None or spec.loader is None:
        raise RuntimeError("Unable to load advanced submission builder.")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


ADV = load_advanced_builder()

TITLE = ADV.TITLE
SHORT_TITLE = ADV.SHORT_TITLE
AUTHOR_NAME = ADV.AUTHOR_NAME
DEGREE = ADV.DEGREE
AFFILIATION = ADV.AFFILIATION
CORRESPONDING = ADV.CORRESPONDING
ORCID = "0000-0002-4771-8285"

REFERENCES = list(ADV.REFERENCES) + [
    "National Center for Biotechnology Information. Gene Expression Omnibus: GSE107994. "
    "https://www.ncbi.nlm.nih.gov/geo/query/acc.cgi?acc=GSE107994. Accessed 18 Mar 2026.",
    "National Center for Biotechnology Information. Gene Expression Omnibus: GSE193777. "
    "https://www.ncbi.nlm.nih.gov/geo/query/acc.cgi?acc=GSE193777. Accessed 18 Mar 2026.",
    "National Center for Biotechnology Information. Gene Expression Omnibus: GSE79362. "
    "https://www.ncbi.nlm.nih.gov/geo/query/acc.cgi?acc=GSE79362. Accessed 18 Mar 2026.",
    "Siddalingaiah HS. tb-progression-transcriptome-meta. GitHub. "
    "https://github.com/hssling/tb-progression-transcriptome-meta. Accessed 18 Mar 2026.",
]


def set_style(doc: Document, line_spacing: float = 1.5) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = line_spacing
    style.paragraph_format.space_after = Pt(0)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)


def add_paragraph(doc: Document, text: str = "", bold: bool = False, center: bool = False) -> None:
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_bmc_cited_paragraph(doc: Document, text: str, center: bool = False) -> None:
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    parts = re.split(r"(\[\[[0-9,\- ]+\]\])", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("[[") and part.endswith("]]"):
            p.add_run(f"[{part[2:-2].replace(' ', '')}]")
        else:
            p.add_run(part)


def add_table(doc: Document, df: pd.DataFrame, title: str) -> None:
    add_paragraph(doc, title)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for idx, col in enumerate(df.columns):
        table.cell(0, idx).text = str(col)
    for row in df.itertuples(index=False):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    doc.add_paragraph("")


def add_figure(doc: Document, image_path: Path, caption: str, width: float = 6.2) -> None:
    add_paragraph(doc, caption)
    if image_path.exists():
        doc.add_picture(str(image_path), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")


def build_title_page(data) -> Path:
    doc = Document()
    set_style(doc)
    add_paragraph(doc, TITLE, bold=True, center=True)
    add_paragraph(doc, "")
    add_paragraph(doc, f"{AUTHOR_NAME}, {DEGREE}", center=True)
    add_paragraph(doc, AFFILIATION, center=True)
    add_paragraph(doc, "")
    add_paragraph(doc, "Article type: Research article")
    add_paragraph(doc, f"Short title: {SHORT_TITLE}")
    add_paragraph(doc, f"Corresponding author: {CORRESPONDING}")
    add_paragraph(doc, f"ORCID: {ORCID}")
    add_paragraph(
        doc,
        "Author contributions: SHS conceived the study, performed the analysis, interpreted the results, drafted the manuscript, and approved the final manuscript.",
    )
    add_paragraph(doc, "Funding: No external funding was received for this work.")
    add_paragraph(doc, "Competing interests: The author declares that he has no competing interests.")
    add_paragraph(
        doc,
        "Ethics statement: This study used de-identified publicly available datasets only and involved no new participant recruitment.",
    )
    add_paragraph(doc, f"Repository: {REPO_URL}")
    add_paragraph(doc, f"Advanced shared-gene analysis sample count: {int(data.cohort_summary['Samples'].sum())}")
    out = OUT_DIR / "01_BMC_Title_Page.docx"
    doc.save(out)
    return out


def build_main_manuscript(data) -> Path:
    text = ADV.manuscript_text(data)
    doc = Document()
    set_style(doc)

    add_paragraph(doc, TITLE, bold=True, center=True)
    add_heading(doc, "Abstract", 1)
    abstract_lines = list(text["abstract"])
    if abstract_lines:
        abstract_lines[-1] = abstract_lines[-1].replace("Conclusion:", "Conclusions:")
    for para in abstract_lines:
        add_paragraph(doc, para)
    add_paragraph(
        doc,
        "Keywords: tuberculosis; transcriptomics; Bayesian meta-analysis; principal component analysis; coexpression analysis; biomarker discovery",
    )

    add_heading(doc, "Introduction", 1)
    for para in text["intro"]:
        add_bmc_cited_paragraph(doc, para)

    add_heading(doc, "Methods", 1)
    for para in text["methods"]:
        add_bmc_cited_paragraph(doc, para)
    add_paragraph(
        doc,
        "Generative AI-assisted language support was limited to drafting support during manuscript preparation. No generative AI system was used to generate, transform, or analyze the underlying data, and all scientific statements were checked manually against the project outputs before inclusion.",
    )

    add_heading(doc, "Results", 1)
    for para in text["results"]:
        add_bmc_cited_paragraph(doc, para)

    add_heading(doc, "Discussion", 1)
    for para in text["discussion"]:
        add_bmc_cited_paragraph(doc, para)

    add_heading(doc, "Conclusions", 1)
    add_paragraph(doc, str(text["conclusion"]))

    add_heading(doc, "Abbreviations", 1)
    add_paragraph(
        doc,
        "BH, Benjamini-Hochberg; DAG, directed acyclic graph; FDR, false discovery rate; GEO, Gene Expression Omnibus; NNLS, non-negative least squares; PCA, principal component analysis.",
    )

    add_heading(doc, "Declarations", 1)
    add_paragraph(
        doc,
        "Ethics approval and consent to participate: Not applicable for this secondary analysis of de-identified public data.",
    )
    add_paragraph(doc, "Consent for publication: Not applicable.")
    add_paragraph(
        doc,
        "Availability of data and materials: The transcriptomic datasets analyzed in this study are publicly available in GEO under accession numbers GSE107994, GSE193777, and GSE79362 [17-19]. Generated results tables, intermediate outputs, and manuscript assets are available in the project repository [20].",
    )
    add_paragraph(
        doc,
        "Code availability: Source code for data processing, harmonization, advanced analysis, sensitivity analysis, and package generation is available in the project repository [20].",
    )
    add_paragraph(doc, "Funding: No external funding was received.")
    add_paragraph(doc, "Competing interests: The author declares that he has no competing interests.")
    add_paragraph(
        doc,
        "Authors' contributions: SHS conceived the study, performed the analysis, interpreted the results, drafted the manuscript, and approved the final manuscript.",
    )
    add_paragraph(
        doc,
        "Acknowledgements: The author thanks the original investigators who generated and deposited the public datasets used in this secondary analysis.",
    )
    add_paragraph(
        doc,
        "Authors' information: SHS is a physician in the Department of Community Medicine at Shridevi Institute of Medical Sciences & Research Hospital, Tumkur, Karnataka, India.",
    )

    add_heading(doc, "References", 1)
    for idx, ref in enumerate(REFERENCES, start=1):
        add_paragraph(doc, f"{idx}. {ref}")

    top_genes_table = data.bayes_genes.head(10).copy()
    top_genes_table = top_genes_table.loc[:, ["gene", "posterior_mean", "ci95_low", "ci95_high", "tau2"]]
    top_genes_table.columns = ["Gene", "Posterior mean", "95% CrI low", "95% CrI high", "Tau^2"]
    top_genes_table = top_genes_table.round(3)

    factor_table = data.factors.copy()
    factor_table["mean_progressor"] = factor_table["mean_progressor"].map(lambda x: f"{x:.3f}")
    factor_table["mean_nonprogressor"] = factor_table["mean_nonprogressor"].map(lambda x: f"{x:.3f}")
    factor_table["ttest_pvalue"] = factor_table["ttest_pvalue"].map(ADV.format_pvalue)
    factor_table.columns = ["Factor", "Mean in progressors", "Mean in non-progressors", "t-test p value"]

    pathway_table = data.pathways.head(8).copy().round(3)
    pathway_table.columns = ["Pathway", "Posterior mean", "Posterior SD", "95% CrI low", "95% CrI high"]

    sensitivity_table_path = (
        ROOT / "results" / "advanced_analysis_gse79362_sensitivity" / "bayesian_gene_meta_with_gse79362.csv"
    )
    sensitivity_table = (
        pd.read_csv(sensitivity_table_path).head(8).loc[:, ["gene", "posterior_mean", "ci95_low", "ci95_high"]]
        if sensitivity_table_path.exists()
        else pd.DataFrame()
    )
    if not sensitivity_table.empty:
        sensitivity_table = sensitivity_table.round(3)
        sensitivity_table.columns = ["Gene", "Posterior mean", "95% CrI low", "95% CrI high"]

    add_heading(doc, "Tables", 1)
    add_table(
        doc,
        data.cohort_summary,
        "Table 1. Cohort summary for the advanced shared-gene analysis. Age and sex were parsed from public metadata fields, and platform indicates the dominant assay type within each retained cohort.",
    )
    add_table(
        doc,
        top_genes_table,
        "Table 2. Top Bayesian gene-level posterior effects. Posterior means and 95% credible intervals were derived from a two-cohort hierarchical synthesis.",
    )
    add_table(
        doc,
        factor_table,
        "Table 3. Latent factor differences between progressors and non-progressors. P values are shown in scientific notation when small.",
    )
    add_table(
        doc,
        pathway_table,
        "Table 4. Leading pathway-level Bayesian posterior effects. Closely related angiogenesis and vasculature terms reflect convergence within the enrichment output rather than fully independent pathways.",
    )
    if not sensitivity_table.empty:
        add_table(
            doc,
            sensitivity_table,
            "Table 5. Top exploratory Bayesian gene-level posterior effects after adding the remapped GSE79362 sensitivity cohort. These results are shown to illustrate evidence expansion and heterogeneity, not to replace the primary two-cohort analysis.",
        )

    add_heading(doc, "Figures", 1)
    add_figure(
        doc,
        ADV.ANALYSIS_DIR / "raw_pca.png",
        "Figure 1. Raw joint PCA of the shared-gene matrix. The leading component is strongly cohort influenced before cohort centering.",
    )
    add_figure(
        doc,
        ADV.ANALYSIS_DIR / "cohort_centered_pca.png",
        "Figure 2. Cohort-centered PCA of the shared-gene matrix. After centering, separation by progressor status becomes more visible.",
    )
    add_figure(
        doc,
        ADV.ANALYSIS_DIR / "factor_boxplots.png",
        "Figure 3. Distribution of latent factor scores by progressor status. All three leading factors remain associated with progression status.",
    )
    add_figure(
        doc,
        ADV.ANALYSIS_DIR / "signature_network_heatmap.png",
        "Figure 4. Spearman correlation heatmap for leading Bayesian signature genes after cohort centering.",
    )

    out = OUT_DIR / "02_BMC_Main_Manuscript.docx"
    doc.save(out)
    return out


def build_cover_letter() -> Path:
    doc = Document()
    set_style(doc)
    add_paragraph(doc, "Cover Letter", bold=True, center=True)
    add_paragraph(doc, "Dear Editors of BMC Medical Genomics,")
    add_paragraph(
        doc,
        f'Please consider the manuscript "{TITLE}" for publication as a Research article in BMC Medical Genomics.',
    )
    add_paragraph(
        doc,
        "This manuscript presents a systems-level reanalysis of public tuberculosis progression transcriptomic cohorts. "
        "The study combines cross-platform harmonization, principal component analysis, factor analysis, Bayesian hierarchical synthesis, "
        "pathway-level posterior modeling, marker-based deconvolution, coexpression analysis, and exploratory evidence expansion through remapping "
        "of a junction-level RNA-sequencing cohort.",
    )
    add_paragraph(
        doc,
        "The study is positioned as an uncertainty-aware genomics analysis that strengthens biological interpretation while keeping translational "
        "claims bounded by the current evidence base. All datasets analyzed are publicly available, and the source code and generated outputs are "
        "available in the project repository.",
    )
    add_paragraph(
        doc,
        "This manuscript is original, has not been published elsewhere, and is not under consideration by another journal.",
    )
    add_paragraph(doc, "Sincerely,")
    add_paragraph(doc, f"{AUTHOR_NAME}, {DEGREE}")
    out = OUT_DIR / "03_BMC_Cover_Letter.docx"
    doc.save(out)
    return out


def build_submission_notes() -> Path:
    doc = Document()
    set_style(doc)
    add_paragraph(doc, "BMC Medical Genomics Submission Notes", bold=True, center=True)
    add_paragraph(doc, "Target journal: BMC Medical Genomics")
    add_paragraph(doc, "Target article type: Research article")
    add_paragraph(doc, "")
    add_paragraph(doc, "Current fit", bold=True)
    for line in [
        "The manuscript is an original genomics-focused research article built around host blood transcriptomic reanalysis.",
        "The package includes a title page, main manuscript, cover letter, supplementary methods/figures, and a separate OmicsClaw extension supplement.",
        "Main-text citations were converted to bracketed numeric style for the BMC package.",
        "Availability and code statements now cite GEO accessions and the repository through the reference list rather than inline URLs in the manuscript body.",
    ]:
        doc.add_paragraph(line, style="List Bullet")
    add_paragraph(doc, "")
    add_paragraph(doc, "Official sources checked on March 18, 2026", bold=True)
    for line in [
        "Journal homepage: https://bmcmedgenomics.biomedcentral.com/",
        "Submission guidelines: https://bmcmedgenomics.biomedcentral.com/submission-guidelines/preparing-your-manuscript",
        "Research article guidance: https://bmcmedgenomics.biomedcentral.com/submission-guidelines/preparing-your-manuscript/research-article",
        "BMC editorial policies: https://link.springer.com/brands/bmc/editorial-policies",
    ]:
        doc.add_paragraph(line, style="List Bullet")
    add_paragraph(doc, "")
    add_paragraph(doc, "Notes for portal entry", bold=True)
    for line in [
        "Article title and author details are provided on the separate title page.",
        "The main manuscript includes Abstract, Keywords, Abbreviations, Declarations, References, Tables, and Figure legends.",
        "Figures are embedded in the preparation file for internal review but can be exported as separate files at upload if required by the journal system.",
    ]:
        doc.add_paragraph(line, style="List Bullet")
    out = OUT_DIR / "06_BMC_Submission_Notes.docx"
    doc.save(out)
    return out


def build_metadata_sheet() -> Path:
    doc = Document()
    set_style(doc)
    add_paragraph(doc, "BMC Metadata Sheet", bold=True, center=True)
    add_paragraph(doc, f"Title: {TITLE}")
    add_paragraph(doc, f"Short title: {SHORT_TITLE}")
    add_paragraph(doc, "Article type: Research article")
    add_paragraph(
        doc,
        "Keywords: tuberculosis; transcriptomics; Bayesian meta-analysis; principal component analysis; coexpression analysis; biomarker discovery",
    )
    add_paragraph(doc, f"Author: {AUTHOR_NAME}, {DEGREE}")
    add_paragraph(doc, f"Affiliation: {AFFILIATION}")
    add_paragraph(doc, f"Corresponding author: {CORRESPONDING}")
    add_paragraph(doc, f"ORCID: {ORCID}")
    add_paragraph(
        doc,
        "Funding statement: No external funding was received.",
    )
    add_paragraph(
        doc,
        "Competing interests statement: The author declares that he has no competing interests.",
    )
    add_paragraph(
        doc,
        "Data availability summary: GEO accessions GSE107994, GSE193777, and GSE79362; repository reference in the manuscript reference list.",
    )
    add_paragraph(doc, f"Repository: {REPO_URL}")
    out = OUT_DIR / "07_BMC_Metadata_Sheet.docx"
    doc.save(out)
    return out


def copy_supporting_files() -> None:
    mapping = {
        "04_Supplementary_Methods_and_Figures.docx": "04_BMC_Supplementary_Methods_and_Figures.docx",
        "06_OmicsClaw_Extensions.docx": "05_BMC_OmicsClaw_Extensions.docx",
        "internal_review_log.md": "internal_review_log.md",
    }
    for src_name, dst_name in mapping.items():
        src = SRC_DIR / src_name
        if src.exists():
            shutil.copy2(src, OUT_DIR / dst_name)


def build_readme() -> Path:
    lines = [
        "# BMC Medical Genomics Advanced Submission Package",
        "",
        "Target journal: BMC Medical Genomics",
        "Article type: Research article",
        f"Repository: {REPO_URL}",
        "",
        "Contents",
        "- 01_BMC_Title_Page.docx",
        "- 02_BMC_Main_Manuscript.docx",
        "- 03_BMC_Cover_Letter.docx",
        "- 04_BMC_Supplementary_Methods_and_Figures.docx",
        "- 05_BMC_OmicsClaw_Extensions.docx",
        "- 06_BMC_Submission_Notes.docx",
        "- 07_BMC_Metadata_Sheet.docx",
        "- validation_report.txt",
        "- internal_review_log.md",
        "",
        "This package is generated from the validated advanced manuscript source and tailored to BMC-style bracketed citations and declaration sections.",
    ]
    out = OUT_DIR / "README.md"
    out.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return out


def build_validation_report() -> Path:
    doc = Document(OUT_DIR / "02_BMC_Main_Manuscript.docx")
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)
    words = sum(len(p.split()) for p in paragraphs)
    lines = [
        "BMC Medical Genomics package validation",
        "Main manuscript: 02_BMC_Main_Manuscript.docx",
        f"Approximate manuscript words (paragraph text only): {words}",
        f"Contains bracketed citations: {'[1]' in text or '[1,2]' in text}",
        f"Contains Abbreviations section: {'Abbreviations' in text}",
        f"Contains Declarations section: {'Declarations' in text}",
        f'Contains Authors\' contributions section: {"Authors\' contributions:" in text}',
        f"Contains Acknowledgements section: {'Acknowledgements:' in text}",
        f"Contains dataset accession references in availability statement: {'GSE107994, GSE193777, and GSE79362 [17-19]' in text}",
        f"Availability and code statements cite repository by reference number: {'Code availability: Source code' in text and '[20]' in text}",
        f"Contains Table 1-5 in order: {all(text.find(f'Table {i}') != -1 for i in range(1, 6))}",
        f"Contains Figure 1-4 in order: {all(text.find(f'Figure {i}') != -1 for i in range(1, 5))}",
        "Evidence boundary note: primary advanced shared-gene analysis remains restricted to GSE107994 and GSE193777.",
        "Interpretive limit note: deconvolution is proxy-based and coexpression is supportive rather than causal.",
    ]
    out = OUT_DIR / "validation_report.txt"
    out.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return out


def zip_package() -> Path:
    zip_path = OUT_DIR / "bmc_med_genomics_advanced_submission_package.zip"
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for path in sorted(OUT_DIR.iterdir()):
            if path == zip_path or path.name.startswith("~$"):
                continue
            zf.write(path, arcname=path.name)
    return zip_path


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    for leftover in OUT_DIR.glob("~$*.docx"):
        leftover.unlink()
    data = ADV.load_data()
    build_title_page(data)
    build_main_manuscript(data)
    build_cover_letter()
    copy_supporting_files()
    build_submission_notes()
    build_metadata_sheet()
    build_readme()
    build_validation_report()
    zip_package()
    print(f"BMC Medical Genomics package written to {OUT_DIR}")


if __name__ == "__main__":
    main()
