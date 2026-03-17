from __future__ import annotations

from pathlib import Path
import re

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
PKG = ROOT / "submission_ready" / "ijtb_20260317_rev3"
MANUSCRIPT = PKG / "02_IJTB_Blinded_Manuscript.docx"
TITLE_PAGE = PKG / "01_IJTB_Title_Page.docx"


def collect_text(doc: Document) -> list[str]:
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


def count_words(text: str) -> int:
    return len([w for w in re.split(r"\s+", text) if w])


def manuscript_metrics() -> dict[str, object]:
    doc = Document(MANUSCRIPT)
    texts = collect_text(doc)
    abstract_start = texts.index("Structured Abstract")
    intro_start = texts.index("Introduction")
    refs_start = texts.index("References")
    tables_start = texts.index("Tables")
    figures_start = texts.index("Figures")
    abstract_text = " ".join(texts[abstract_start + 1 : intro_start])
    body_text = " ".join(texts[intro_start:refs_start])
    refs = [t for t in texts[refs_start + 1 : tables_start] if re.match(r"^\d+\.", t)]
    table_titles = [t for t in texts[tables_start + 1 : figures_start] if t.startswith("Table ")]
    figure_titles = [t for t in texts[figures_start + 1 :] if t.startswith("Fig. ")]

    superscripts = []
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.superscript and run.text.strip():
                superscripts.append(run.text.strip())

    seq = []
    for item in superscripts:
        for part in re.split(r"[,\-]", item):
            if part.isdigit():
                seq.append(int(part))

    return {
        "abstract_words": count_words(abstract_text),
        "body_words": count_words(body_text),
        "reference_count": len(refs),
        "table_count": len(table_titles),
        "figure_count": len(figure_titles),
        "superscript_count": len(superscripts),
        "citation_min": min(seq) if seq else None,
        "citation_max": max(seq) if seq else None,
        "has_author_name_in_main": "Siddalingaiah" in " ".join(texts[:40]),
    }


def title_page_metrics() -> dict[str, object]:
    doc = Document(TITLE_PAGE)
    text = " ".join(collect_text(doc))
    return {
        "has_short_title": "Short title:" in text,
        "has_corresponding": "Corresponding author:" in text,
        "has_conflict": "Conflict of interest:" in text,
    }


def main() -> None:
    m = manuscript_metrics()
    t = title_page_metrics()
    report = [
        f"abstract_words={m['abstract_words']}",
        f"body_words={m['body_words']}",
        f"reference_count={m['reference_count']}",
        f"table_count={m['table_count']}",
        f"figure_count={m['figure_count']}",
        f"superscript_count={m['superscript_count']}",
        f"citation_min={m['citation_min']}",
        f"citation_max={m['citation_max']}",
        f"has_author_name_in_main={m['has_author_name_in_main']}",
        f"title_has_short_title={t['has_short_title']}",
        f"title_has_corresponding={t['has_corresponding']}",
        f"title_has_conflict={t['has_conflict']}",
    ]
    out = PKG / "ijtb_validation_report.txt"
    out.write_text("\n".join(report) + "\n", encoding="utf-8")
    print(out.relative_to(ROOT))
    print("\n".join(report))


if __name__ == "__main__":
    main()
