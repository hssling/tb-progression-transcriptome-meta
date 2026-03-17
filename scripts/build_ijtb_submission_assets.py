from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re
import zipfile

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "submission_ready" / "ijtb_20260317_rev3"
REPO_URL = "https://github.com/hssling/tb-progression-transcriptome-meta"

AUTHOR_NAME = "Siddalingaiah H S"
AFFILIATION = "Department of Community Medicine, Shridevi Institute of Medical Sciences & Research Hospital, Tumkur, Karnataka, India"
CORRESPONDING = "Siddalingaiah H S, Department of Community Medicine, Shridevi Institute of Medical Sciences & Research Hospital, Tumkur, Karnataka, India. Email: hssling@yahoo.com"
SHORT_TITLE = "TB progression transcriptomics"
TITLE = "Host blood transcriptomic signatures associated with tuberculosis progression: a reproducible public-data meta-analysis"

REFERENCE_LIST = [
    "World Health Organization. Global tuberculosis report 2025. Geneva: World Health Organization; 2025.",
    "Houben RMGJ, Dodd PJ. The global burden of latent tuberculosis infection: a re-estimation using mathematical modelling. PLoS Med. 2016;13:e1002152.",
    "Zak DE, Penn-Nicholson A, Scriba TJ, et al. A blood RNA signature for tuberculosis disease risk: a prospective cohort study. Lancet. 2016;387:2312-2322.",
    "Singhania A, Verma R, Graham CM, et al. A modular transcriptional signature identifies phenotypic heterogeneity of human tuberculosis infection. Nat Commun. 2018;9:2308.",
    "Rajamanickam A, Munisankar S, Dolla CK, et al. Host blood-based biosignatures for subclinical TB and incipient TB: a prospective study of adult TB household contacts in Southern India. Front Immunol. 2022;13:1065779.",
    "Sweeney TE, Braviak L, Tato CM, Khatri P. Genome-wide expression for diagnosis of pulmonary tuberculosis: a multicohort analysis. Lancet Respir Med. 2016;4:213-224.",
    "Warsinske HC, Rao AM, Moreira FMF, et al. Assessment of validity of a blood-based 3-gene signature score for progression and diagnosis of tuberculosis, disease severity, and treatment response. JAMA Netw Open. 2018;1:e183779.",
    "Davis S, Meltzer PS. GEOquery: a bridge between the Gene Expression Omnibus (GEO) and BioConductor. Bioinformatics. 2007;23:1846-1847.",
    "Ritchie ME, Phipson B, Wu D, et al. limma powers differential expression analyses for RNA-sequencing and microarray studies. Nucleic Acids Res. 2015;43:e47.",
    "Leek JT, Johnson WE, Parker HS, Jaffe AE, Storey JD. The sva package for removing batch effects and other unwanted variation in high-throughput experiments. Bioinformatics. 2012;28:882-883.",
    "Balduzzi S, Rucker G, Schwarzer G. How to perform a meta-analysis with R: a practical tutorial. Evid Based Ment Health. 2019;22:153-160.",
    "Friedman J, Hastie T, Tibshirani R. Regularization paths for generalized linear models via coordinate descent. J Stat Softw. 2010;33:1-22.",
    "Wu T, Hu E, Xu S, et al. clusterProfiler 4.0: a universal enrichment tool for interpreting omics data. Innovation (Camb). 2021;2:100141.",
    "Wesche H, Gao X, Li X, et al. IRAK-M is a novel member of the interleukin-1 receptor-associated kinase family. J Biol Chem. 1999;274:19403-19410.",
    "tb-progression-transcriptome-meta repository. Available from: https://github.com/hssling/tb-progression-transcriptome-meta.",
]


@dataclass
class ManuscriptData:
    dataset_summary: pd.DataFrame
    meta_gene_list: pd.DataFrame
    loco_performance: pd.DataFrame
    validation_summary: pd.DataFrame
    pathway_enrichment: pd.DataFrame
    lasso_signature: pd.DataFrame


def set_style(doc: Document, double_spacing: bool = True) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 2 if double_spacing else 1.15
    style.paragraph_format.space_after = Pt(0)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_plain_paragraph(doc: Document, text: str = "", bold: bool = False, center: bool = False) -> None:
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold


def add_cited_paragraph(doc: Document, text: str, center: bool = False) -> None:
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    parts = re.split(r"(\[\[[0-9,\- ]+\]\])", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("[[") and part.endswith("]]"):
            run = p.add_run(part[2:-2].replace(" ", ""))
            run.font.superscript = True
        else:
            p.add_run(part)


def add_table(doc: Document, df: pd.DataFrame, title: str, round_map: dict[str, int] | None = None) -> None:
    add_plain_paragraph(doc, title)
    table_df = df.copy()
    if round_map:
        for col, digits in round_map.items():
            if col in table_df.columns:
                table_df[col] = table_df[col].map(lambda v: f"{v:.{digits}f}" if pd.notna(v) and isinstance(v, (int, float)) else str(v))
    table = doc.add_table(rows=1, cols=len(table_df.columns))
    table.style = "Table Grid"
    for idx, col in enumerate(table_df.columns):
        table.cell(0, idx).text = str(col)
    for row in table_df.itertuples(index=False):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    doc.add_paragraph("")


def add_figure(doc: Document, image_path: Path, caption: str, width: float = 6.0) -> None:
    add_plain_paragraph(doc, caption)
    if image_path.exists():
        doc.add_picture(str(image_path), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")


def load_data() -> ManuscriptData:
    return ManuscriptData(
        dataset_summary=pd.read_csv(ROOT / "R_pipeline" / "output" / "dataset_summary.csv"),
        meta_gene_list=pd.read_csv(ROOT / "submission_package" / "Tables" / "meta_gene_list.csv"),
        loco_performance=pd.read_csv(ROOT / "results" / "tables" / "loco_performance.csv"),
        validation_summary=pd.read_csv(ROOT / "R_pipeline" / "output" / "validation_summary.csv"),
        pathway_enrichment=pd.read_csv(ROOT / "R_pipeline" / "output" / "pathway_enrichment_R.csv"),
        lasso_signature=pd.read_csv(ROOT / "R_pipeline" / "output" / "lasso_signature.csv"),
    )


def build_title_page() -> Path:
    doc = Document()
    set_style(doc)
    add_plain_paragraph(doc, TITLE, bold=True, center=True)
    add_plain_paragraph(doc, "")
    add_plain_paragraph(doc, AUTHOR_NAME, center=True)
    add_plain_paragraph(doc, AFFILIATION, center=True)
    add_plain_paragraph(doc, "")
    add_plain_paragraph(doc, f"Short title: {SHORT_TITLE}")
    add_plain_paragraph(doc, f"Corresponding author: {CORRESPONDING}")
    add_plain_paragraph(doc, "Author ORCID: 0000-0002-4771-8285")
    add_plain_paragraph(doc, "Author contributions: Single-author study. Siddalingaiah H S conceived the study, supervised analysis interpretation, prepared the manuscript, and approved the final version.")
    add_plain_paragraph(doc, "Funding: No external funding was received for this work.")
    add_plain_paragraph(doc, "Conflict of interest: The author declares no competing interests.")
    add_plain_paragraph(doc, "Ethics statement: This study used de-identified publicly available transcriptomic datasets and involved no new human participant recruitment.")
    out = OUT_DIR / "01_IJTB_Title_Page.docx"
    doc.save(out)
    return out


def build_main_manuscript(data: ManuscriptData) -> Path:
    doc = Document()
    set_style(doc)
    add_plain_paragraph(doc, TITLE, bold=True, center=True)

    add_heading(doc, "Structured Abstract", 1)
    add_plain_paragraph(doc, "Background/Objectives: Tuberculosis remains a major global health problem, and better tools are needed to identify progression from latent or incipient infection to active disease. Blood transcriptomic signatures are promising, but reproducibility across cohorts and platforms remains limited.")
    add_plain_paragraph(doc, "Methods: We developed a reproducible public-data workflow for tuberculosis progression transcriptomics using curated GEO cohorts. Three primary cohorts were processed, and two cohorts with harmonized binary labels were used for leave-one-cohort-out validation. Gene-level effects were synthesized with a random-effects approach, and three classifiers were assessed: elastic net, linear support vector machine, and a gene-set score.")
    add_plain_paragraph(doc, "Results: The analysis prioritized a 25-gene host signature associated with progression. Leading genes included MILR1, VSIG4, CCR2, CD36, FZD5, AQP1, and IRAK3. The best validation result was produced by the gene-set-score model in the left-out GSE107994 cohort, with AUC-ROC 0.914 and AUC-PR 0.828. Mean AUC-ROC values were 0.884 for the gene-set-score model, 0.838 for linear support vector machine, and 0.780 for elastic net. Enrichment analysis highlighted phagocytosis, myeloid activation, immune regulation, and vascular remodeling.")
    add_plain_paragraph(doc, "Conclusions: This analysis identified a 25-gene host signature led by MILR1, VSIG4, CCR2, CD36, FZD5, AQP1, and IRAK3, with best leave-one-cohort-out AUC-ROC 0.914. The pathway profile supports phagocytosis, myeloid activation, immune regulation, and vascular remodeling. These findings indicate a reproducible host-response signal linked to tuberculosis progression, but further external validation is needed before clinical use.")
    add_plain_paragraph(doc, "Keywords: biomarker; host response; latent tuberculosis infection; machine learning; transcriptomics")

    add_heading(doc, "Introduction", 1)
    intro = [
        "Tuberculosis continues to cause substantial morbidity and mortality worldwide. The World Health Organization estimated that 10.7 million people developed tuberculosis in 2024 and 1.23 million died from the disease, underscoring the need for earlier and more accurate risk stratification.[[1]]",
        "Most infected individuals do not progress to active disease, yet the population burden of latent infection remains very large. This creates a practical need for tools that can distinguish people who are likely to progress from those who remain clinically stable.[[2]]",
        "Blood-based host transcriptomic signatures are attractive because they are non-sputum based, mechanistically informative, and potentially useful in asymptomatic high-risk contacts. Important studies have already shown that blood RNA signatures can anticipate disease risk and capture biological heterogeneity along the tuberculosis spectrum.[[3-6]]",
        "However, reproducibility remains the main obstacle. Published signatures often lose performance when applied to independent cohorts with different recruitment designs, assay platforms, or preprocessing pipelines. Multi-cohort evaluations have also shown that transcriptomic classifiers are sensitive to hidden study structure and analytic choices.[[6,7]]",
        "Public datasets are therefore useful only if they are handled conservatively. A scientifically sound public-data study must separate cohort discovery from cohort validation, avoid information leakage, and keep its claims bounded by the actual harmonized labels that are available.",
        "The present study was designed as a reproducible public-data meta-analysis aimed at identifying host blood genes associated with tuberculosis progression. The secondary aim was to package the analysis as a journal-ready submission while retaining transparent reporting, simple language, and strict alignment between the manuscript claims and the available outputs."
    ]
    for paragraph in intro:
        add_cited_paragraph(doc, paragraph)

    add_heading(doc, "Methods", 1)
    methods = {
        "Study design": "This was a secondary analysis of publicly available transcriptomic cohorts relevant to tuberculosis progression, incipient tuberculosis, or closely related host-response states. The analysis was organized as a reproducible workflow that included cohort curation, harmonization, random-effects synthesis, pathway interpretation, and cross-cohort validation.",
        "Data sources and cohort selection": "Curated cohorts were identified from GEO using progression-focused tuberculosis terms. The current submission retains three primary processed cohorts: GSE107994, GSE193777, and GSE79362. Of these, two cohorts had directly comparable binary labels suitable for leave-one-cohort-out validation. Additional microarray cohorts were retained as candidate resources for future expansion rather than being forced into the present validation framework.[[3-8]]",
        "Preprocessing and harmonization": "Expression matrices were normalized and mapped to gene-level features wherever possible. Duplicate or ambiguous mappings were collapsed conservatively, and datasets were restricted to a shared feature space before downstream analysis. This conservative approach was chosen to reduce false precision and to keep the manuscript aligned with the exported outputs.[[8,9]]",
        "Meta-analysis": "Per-gene effects for progressor versus non-progressor comparisons were combined using a random-effects framework. Genes were ranked by absolute meta z score, and false-discovery-adjusted values were retained for prioritization. The synthesis was used to identify robust candidate genes rather than to make exaggerated claims of universal invariance.[[10]]",
        "Model development and validation": "Three classifier types were assessed: elastic net logistic regression, linear support vector machine, and a simple gene-set score based on the ranked signature. Validation used leave-one-cohort-out testing so that one complete cohort was excluded from model training and used only for evaluation. This design was chosen specifically to reduce information leakage and to reflect realistic cross-cohort transfer.[[11,12]]",
        "Functional interpretation": "The ranked genes were interpreted using exported Gene Ontology biological-process enrichment results generated by the reporting pipeline. Only pathways present in the exported results were discussed in the manuscript, which limited over-interpretation and improved traceability.[[13]]",
        "Ethics": "The present study used de-identified publicly available data only. No new patients were recruited, and no local intervention or specimen collection was performed. Ethical approvals and informed consent for the original cohorts were handled by the primary investigators as described in the original publications."
    }
    for heading, paragraph in methods.items():
        add_heading(doc, heading, 2)
        add_cited_paragraph(doc, paragraph)

    add_heading(doc, "Results", 1)
    add_heading(doc, "Cohort characteristics", 2)
    add_cited_paragraph(doc, "The curated registry identified six relevant public studies, of which three were processed as primary cohorts in the current package. These cohorts represented the Leicester cohort, a Southern India household-contact cohort, and the ACS prospective cohort. Together they accounted for 533 processed samples and captured both discovery-rich and validation-ready study structures (Table 1).")
    add_cited_paragraph(doc, "The processed cohorts were not fully interchangeable. Only two cohorts provided directly comparable binary labels for the main leave-one-cohort-out analysis. That limitation defines the real boundary of the current evidence and helps prevent overstatement of model maturity.")
    add_cited_paragraph(doc, "The cohort mix nevertheless remains informative because it reflects the actual structure of the public tuberculosis transcriptomic literature. Some studies are rich in biological context but less suitable for direct binary validation, whereas others are better suited to classifier testing. Treating these cohorts as analytically identical would create false precision. The present workflow preserves that distinction and uses it to guide how far the results can reasonably be interpreted.")

    add_heading(doc, "Progression-associated transcriptomic signals", 2)
    add_cited_paragraph(doc, "The random-effects synthesis prioritized genes involved in innate immune regulation, phagocytic handling, leukocyte trafficking, and vascular remodeling. The strongest positive signals were MILR1, VSIG4, CCR2, CD36, FZD5, AQP1, and CRISPLD2, while EPN2 and PLD4 showed strong inverse directionality among the leading ranks (Table 2).")
    add_cited_paragraph(doc, "The top-ranked genes were biologically coherent rather than statistically isolated. MILR1 and IRAK3 suggest regulatory signaling, CCR2 and ITGB2 support immune-cell recruitment, CD36 reflects macrophage-associated lipid and scavenger biology, and FZD5 points to a broader signaling context that may relate to tissue remodeling and inflammatory balance. A forest-style summary of the leading transcriptomic effects is shown in Fig. 1.")
    add_cited_paragraph(doc, "The leading genes also showed a practical balance between effect size and interpretability. For a translational study, it is not enough for genes to rank highly; they should also map to host processes that clinicians and laboratory scientists can understand. In that sense, the current list is more convincing than an opaque machine-learning feature set because it links progression to recognizable pathways in myeloid activation, adhesion, and immune regulation.")

    add_heading(doc, "Cross-cohort validation performance", 2)
    add_cited_paragraph(doc, "The strongest leave-one-cohort-out performance was obtained with the gene-set-score model when GSE107994 was used as the held-out cohort, producing AUC-ROC 0.914 and AUC-PR 0.828. Across the two held-out cohorts, the mean AUC-ROC values were 0.884 for the gene-set-score model, 0.838 for linear support vector machine, and 0.780 for elastic net (Table 3).")
    add_cited_paragraph(doc, "The results also showed that discrimination and calibration were not equivalent. The gene-set-score model ranked samples best, whereas linear support vector machine achieved more favorable Brier scores. This pattern suggests that the host signal is strong enough for cross-cohort discrimination, while absolute risk estimation remains less mature. The ROC comparison is presented in Fig. 2.")
    add_cited_paragraph(doc, "The difference between the held-out cohorts is also informative. A signature that performs identically in every public cohort is uncommon in this field because cohorts differ in baseline risk, recruitment frame, sampling window, and assay processing. The observed variation therefore does not invalidate the signature; instead, it underlines why strict external testing is necessary before any host transcriptomic model is presented as clinically stable.")

    add_heading(doc, "Functional interpretation", 2)
    add_cited_paragraph(doc, "Pathway enrichment emphasized regulation of angiogenesis, vasculature development, immune effector processes, leukocyte degranulation, myeloid activation, and phagocytosis. Recurrent genes such as CCR2, CD36, FZD5, ITGB2, and SPARC connected the pathway-level findings to the ranked signature, supporting a model in which progression risk is associated with coordinated immune trafficking, tissue-interface remodeling, and altered innate regulation.")
    add_cited_paragraph(doc, "IRAK3 deserves specific mention because it is a negative regulator of Toll-like receptor signaling. Its prominence in the ranked genes suggests that progression is not defined only by inflammatory activation, but also by dysregulated or compensatory immune restraint.[[14]]")
    add_cited_paragraph(doc, "This interpretation is clinically useful because it avoids a simplistic view in which progression is reduced to 'more inflammation'. The enrichment profile instead suggests a more complex biological state in which recruitment, regulation, phagocytic behavior, and tissue-level adaptation change together. That framing is easier to defend scientifically and easier to communicate to a broad tuberculosis readership.")

    add_heading(doc, "Discussion", 1)
    discussion = [
        "This study shows that a coherent host blood transcriptomic program associated with tuberculosis progression can be recovered from public cohorts when the workflow is transparent and validation is performed at the cohort level. The value of the present work lies not only in the ranked genes but also in the disciplined handling of evidence boundaries.",
        "The leading genes form a plausible biological module. MILR1, VSIG4, CCR2, CD36, FZD5, AQP1, and IRAK3 together suggest a state characterized by myeloid activity, cellular recruitment, phagocytic remodeling, and regulatory feedback. This picture is biologically more credible than a purely generic inflammatory signal because the genes point to related host processes rather than unrelated extremes.",
        "The findings are also consistent with the broader tuberculosis progression literature. Zak et al. established that blood RNA could anticipate disease risk, Singhania et al. emphasized transcriptional heterogeneity, and Rajamanickam et al. extended the concept to subclinical and incipient tuberculosis in household contacts.[[3-5]] The present analysis adds value by prioritizing reproducibility, clear reporting, and cohort-held-out validation.",
        "A major strength of the current workflow is its resistance to information leakage. Multi-cohort transcriptomic studies can appear stronger than they really are when data from the same cohort contribute to both model development and model testing. The use of leave-one-cohort-out validation directly addresses that issue and makes the reported performance more credible, even if it appears less flattering than row-wise split results.",
        "The study also has important limitations. Only two cohorts presently support the harmonized binary validation framework used for the main predictive claims. Additional candidate microarray cohorts remain outside the finalized validation path, and the current models show better ranking discrimination than calibration. These points mean that the signature should be interpreted as a candidate for further validation rather than an immediately deployable clinical test.",
        "Another limitation is that public metadata do not always preserve the time-to-event resolution needed for stronger statements about how many months before disease onset the signal becomes clinically useful. The current manuscript therefore avoids claims that go beyond the available labels.",
        "The balance between simplicity and robustness is another important point. The gene-set-score model performed well despite its relative simplicity, which is encouraging from a translational perspective. Complex models can be attractive computationally, but simpler score-based approaches are often easier to audit, explain, and translate into low-resource laboratory settings. The present results suggest that parsimony may be an advantage rather than a weakness in this domain.",
        "The signature also remains clinically interpretable. Genes such as CCR2, CD36, VSIG4, ITGB2, and IRAK3 do not read like arbitrary machine-selected identifiers; they point toward host processes that tuberculosis clinicians and immunology researchers already recognize. That interpretability matters because biomarker adoption depends not only on numerical performance but also on whether the biology appears credible and understandable.",
        "Even with these limitations, the study remains useful. It identifies a biologically interpretable candidate gene set, demonstrates that cross-cohort signal retention is achievable, and provides a reproducible submission package that can be updated as more progression-ready cohorts are integrated.",
        "Future work should focus on three areas: adding more prospective progression cohorts, improving calibration after broader cohort harmonization, and translating the most stable genes into practical multiplex assays. The latter step should prioritize genes that are statistically strong, mechanistically understandable, and technically feasible for targeted testing.",
        "Overall, the manuscript is best read as a rigorous biomarker-discovery and reproducibility study. That framing is scientifically more honest than presenting the current signature as a finished diagnostic tool, and it is more likely to remain useful as additional public datasets become available."
    ]
    for paragraph in discussion:
        add_cited_paragraph(doc, paragraph)

    add_heading(doc, "Conclusion", 1)
    add_cited_paragraph(doc, "A reproducible public-data meta-analysis identified a 25-gene host blood signature associated with tuberculosis progression, with leading contributions from MILR1, VSIG4, CCR2, CD36, FZD5, AQP1, and IRAK3. Cross-cohort validation showed that this signal remained detectable outside the discovery setting, with the strongest held-out performance observed in GSE107994. The pathway profile suggests coordinated phagocytic, myeloid, regulatory, and vascular-remodeling responses. These findings support the presence of a transferable host-response signal, but broader external validation and improved calibration are still required before any clinical application can be justified.")

    add_heading(doc, "Acknowledgements", 1)
    add_plain_paragraph(doc, "The author acknowledges the investigators who generated and shared the public GEO datasets used in this analysis.")
    add_heading(doc, "Funding", 1)
    add_plain_paragraph(doc, "No external funding was received for this work.")
    add_heading(doc, "Conflict of interest", 1)
    add_plain_paragraph(doc, "The author declares that there is no conflict of interest.")
    add_heading(doc, "Ethical approval", 1)
    add_plain_paragraph(doc, "This study was based entirely on de-identified publicly available datasets. No new human participant recruitment or local specimen collection was performed.")
    add_heading(doc, "Consent", 1)
    add_plain_paragraph(doc, "Not applicable to the current secondary analysis. Original cohort-specific ethics and consent procedures were handled by the primary investigators.")
    add_heading(doc, "Data availability", 1)
    add_cited_paragraph(doc, "All generated submission files, analysis scripts, and supporting outputs are available in a public repository.[[15]]")
    add_heading(doc, "Code availability", 1)
    add_cited_paragraph(doc, "The reproducible codebase used for this analysis and document preparation is available in a public repository.[[15]]")
    add_heading(doc, "Declaration on generative AI use", 1)
    add_plain_paragraph(doc, "During manuscript preparation, the author used a generative AI-assisted coding and drafting tool to help structure documents, refine language, and prepare submission files. The author critically reviewed, edited, and accepted full responsibility for the final content.")

    add_heading(doc, "References", 1)
    for idx, ref in enumerate(REFERENCE_LIST, start=1):
        add_plain_paragraph(doc, f"{idx}. {ref}")

    add_heading(doc, "Tables", 1)
    table1 = data.dataset_summary.loc[data.dataset_summary["Status"].str.contains("Primary", na=False), ["GEO_ID", "Platform", "Total_Samples", "Groups", "PMID"]].rename(columns={"GEO_ID": "Cohort", "Total_Samples": "Samples", "Groups": "Clinical groups"})
    add_table(doc, table1, "Table 1. Primary cohorts included in the current submission-ready analysis.")
    table2 = data.meta_gene_list.loc[:, ["gene", "meta_effect", "meta_z", "meta_fdr", "i2"]].head(12).rename(columns={"gene": "Gene", "meta_effect": "Meta effect", "meta_z": "Meta z score", "meta_fdr": "FDR", "i2": "I2"})
    add_table(doc, table2, "Table 2. Leading progression-associated genes from the random-effects synthesis.", {"Meta effect": 3, "Meta z score": 3, "I2": 1})
    table3 = data.loco_performance.rename(columns={"left_out_cohort": "Left-out cohort", "auc_roc": "AUC-ROC", "auc_pr": "AUC-PR", "brier": "Brier score"})
    add_table(doc, table3, "Table 3. Leave-one-cohort-out model performance.", {"AUC-ROC": 3, "AUC-PR": 3, "Brier score": 3})

    add_heading(doc, "Figures", 1)
    fig_dir = ROOT / "R_pipeline" / "figures"
    add_figure(doc, fig_dir / "forest_plot.png", "Fig. 1. Forest-style summary of the leading transcriptomic effects associated with tuberculosis progression.")
    add_figure(doc, fig_dir / "roc_combined.png", "Fig. 2. Receiver operating characteristic curves for the evaluated validation models.")

    out = OUT_DIR / "02_IJTB_Blinded_Manuscript.docx"
    doc.save(out)
    return out


def build_cover_letter() -> Path:
    doc = Document()
    set_style(doc)
    add_plain_paragraph(doc, "17 March 2026")
    add_plain_paragraph(doc, "The Editor")
    add_plain_paragraph(doc, "Indian Journal of Tuberculosis")
    add_plain_paragraph(doc, "")
    body = [
        f'I am submitting the manuscript titled "{TITLE}" for consideration as an Original Article in the Indian Journal of Tuberculosis.',
        "This manuscript presents a reproducible public-data meta-analysis of host blood transcriptomic signatures associated with tuberculosis progression. The work integrates cohort curation, conservative harmonization, random-effects synthesis, leave-one-cohort-out validation, and journal-ready reporting.",
        "The manuscript is scientifically cautious by design. It reports a transferable host-response signal while clearly acknowledging the current evidence limits, including the restricted number of harmonized validation-ready cohorts and the need for further external validation.",
        "The submission is original, has not been published previously, is not under consideration elsewhere, and is based entirely on de-identified publicly available datasets.",
        "All files required for submission, including title page, blinded manuscript, cover letter, author statements, checklist, and an optional supplementary appendix, are included in the current package.",
        f"Repository: {REPO_URL}",
        f"Corresponding author: {CORRESPONDING}",
    ]
    for paragraph in body:
        add_plain_paragraph(doc, paragraph)
    add_plain_paragraph(doc, "")
    add_plain_paragraph(doc, "Sincerely,")
    add_plain_paragraph(doc, AUTHOR_NAME)
    out = OUT_DIR / "03_IJTB_Cover_Letter.docx"
    doc.save(out)
    return out


def build_author_statements() -> Path:
    doc = Document()
    set_style(doc)
    add_heading(doc, "Author Statements", 1)
    statements = {
        "Authorship": "Single-author manuscript. Siddalingaiah H S conceived the study, reviewed the analytical outputs, wrote the manuscript, and approved the final version.",
        "Funding": "No funding was received.",
        "Conflict of interest": "No competing interests are declared.",
        "Ethics": "Secondary analysis of de-identified public data only.",
        "Consent": "Not applicable for the present analysis.",
        "Data availability": f"Repository and generated files are available at {REPO_URL}.",
        "Code availability": f"Code used for analysis and submission-file generation is available at {REPO_URL}.",
        "Generative AI use": "A generative AI-assisted coding and drafting tool was used to help organize and refine manuscript and submission documents. The author critically reviewed all outputs and takes full responsibility for the content.",
    }
    for heading, text in statements.items():
        add_heading(doc, heading, 2)
        add_plain_paragraph(doc, text)
    out = OUT_DIR / "04_IJTB_Author_Statements.docx"
    doc.save(out)
    return out


def build_checklist() -> Path:
    doc = Document()
    set_style(doc)
    add_heading(doc, "IJTB Submission Checklist", 1)
    items = [
        "Separate title page prepared with author details and corresponding author information.",
        "Blinded main manuscript prepared without author identifiers.",
        "Original article structure includes Introduction, Methods, Results, Discussion, and Conclusion.",
        "Structured abstract kept within the journal limit.",
        "References numbered in sequential order and cited in superscript format.",
        "Number of references kept within the journal limit.",
        "Number of tables limited to three.",
        "Number of figures limited to two.",
        "Tables and figures placed at the end of the manuscript after the references section.",
        "Funding, conflict of interest, ethics, consent, data availability, and code availability statements included.",
        "Cover letter prepared.",
        "Highlights file prepared for optional Elsevier submission fields.",
        "Optional supplementary appendix prepared for supporting information not placed in the main article.",
    ]
    for idx, item in enumerate(items, start=1):
        add_plain_paragraph(doc, f"{idx}. {item}")
    out = OUT_DIR / "05_IJTB_Submission_Checklist.docx"
    doc.save(out)
    return out


def build_highlights() -> tuple[Path, Path]:
    bullets = [
        "Public TB transcriptomic cohorts supported reproducible progression-signal discovery.",
        "A 25-gene host signature was led by MILR1, VSIG4, CCR2, CD36, and FZD5.",
        "Best leave-one-cohort-out performance reached AUC-ROC 0.914 in GSE107994.",
        "Biological signals suggested phagocytosis, myeloid activation, and immune regulation.",
        "Current evidence supports biomarker discovery, not immediate clinical deployment.",
    ]
    txt_path = OUT_DIR / "07_IJTB_Highlights.txt"
    txt_path.write_text("\n".join(f"- {bullet}" for bullet in bullets) + "\n", encoding="utf-8")

    doc = Document()
    set_style(doc, double_spacing=False)
    add_heading(doc, "Highlights", 1)
    add_plain_paragraph(doc, "Prepared for optional Elsevier highlights upload.")
    for bullet in bullets:
        add_plain_paragraph(doc, f"- {bullet}")
    docx_path = OUT_DIR / "07_IJTB_Highlights.docx"
    doc.save(docx_path)
    return txt_path, docx_path


def build_supplement(data: ManuscriptData) -> Path:
    doc = Document()
    set_style(doc)
    add_heading(doc, "Supplementary Appendix", 1)
    add_plain_paragraph(doc, "This appendix provides supporting material that was not retained in the main manuscript because of journal limits on tables and figures.")
    full_genes = data.meta_gene_list.loc[:, ["gene", "meta_effect", "meta_z", "meta_fdr", "i2"]].head(25).rename(columns={"gene": "Gene", "meta_effect": "Meta effect", "meta_z": "Meta z score", "meta_fdr": "FDR", "i2": "I2"})
    add_table(doc, full_genes, "Supplementary Table S1. Top 25 progression-associated genes.", {"Meta effect": 3, "Meta z score": 3, "I2": 1})
    enrich = data.pathway_enrichment.loc[:, ["ID", "Description", "GeneRatio", "p.adjust", "geneID"]].head(15).rename(columns={"p.adjust": "Adjusted p value", "geneID": "Genes"})
    add_table(doc, enrich, "Supplementary Table S2. Leading Gene Ontology biological processes.")
    stable = data.lasso_signature.loc[:, ["rank", "gene", "stability", "meta_z"]].head(15).rename(columns={"rank": "Rank", "gene": "Gene", "stability": "Stability", "meta_z": "Meta z score"})
    add_table(doc, stable, "Supplementary Table S3. Stability-ranked minimal signature candidates.", {"Stability": 2, "Meta z score": 3})
    out = OUT_DIR / "06_IJTB_Supplementary_Appendix.docx"
    doc.save(out)
    return out


def build_readme() -> Path:
    text = "\n".join([
        "# IJTB submission package",
        "",
        "Target journal: Indian Journal of Tuberculosis (IJTB)",
        f"Repository: {REPO_URL}",
        "",
        "Files:",
        "- 01_IJTB_Title_Page.docx",
        "- 02_IJTB_Blinded_Manuscript.docx",
        "- 03_IJTB_Cover_Letter.docx",
        "- 04_IJTB_Author_Statements.docx",
        "- 05_IJTB_Submission_Checklist.docx",
        "- 06_IJTB_Supplementary_Appendix.docx",
        "- 07_IJTB_Highlights.txt",
        "- 07_IJTB_Highlights.docx",
        "- ijtb_submission_package.zip",
    ])
    out = OUT_DIR / "README.md"
    out.write_text(text + "\n", encoding="utf-8")
    return out


def package(files: list[Path]) -> Path:
    zip_path = OUT_DIR / "ijtb_submission_package.zip"
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for file in files:
            zf.write(file, arcname=file.name)
    return zip_path


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    data = load_data()
    files = [
        build_title_page(),
        build_main_manuscript(data),
        build_cover_letter(),
        build_author_statements(),
        build_checklist(),
        build_supplement(data),
        build_readme(),
    ]
    txt_highlights, docx_highlights = build_highlights()
    files.extend([txt_highlights, docx_highlights])
    files.append(package(files))
    for file in files:
        print(file.relative_to(ROOT))


if __name__ == "__main__":
    main()
